'''
  ******************************************************************************************
	  Assembly:                Chonky
	  Filename:                processing.py
	  Author:                  Terry D. Eppler
	  Created:                 05-31-2022

	  Last Modified By:        Terry D. Eppler
	  Last Modified On:        05-01-2025
  ******************************************************************************************
  <copyright file="processing.py" company="Terry D. Eppler">

		 Boo is a df analysis tool that integrates various Generative AI, Text-Processing, and
		 Machine-Learning algorithms for federal analysts.
		 Copyright ©  2022  Terry Eppler

	 Permission is hereby granted, free of charge, to any person obtaining a copy
	 of this software and associated documentation files (the “Software”),
	 to deal in the Software without restriction,
	 including without limitation the rights to use,
	 copy, modify, merge, publish, distribute, sublicense,
	 and/or sell copies of the Software,
	 and to permit persons to whom the Software is furnished to do so,
	 subject to the following conditions:

	 The above copyright notice and this permission notice shall be included in all
	 copies or substantial portions of the Software.

	 THE SOFTWARE IS PROVIDED “AS IS”, WITHOUT WARRANTY OF ANY KIND, EXPRESS OR IMPLIED,
	 INCLUDING BUT NOT LIMITED TO THE WARRANTIES OF MERCHANTABILITY,
	 FITNESS FOR A PARTICULAR PURPOSE AND NON-INFRINGEMENT.
	 IN NO EVENT SHALL THE AUTHORS OR COPYRIGHT HOLDERS BE LIABLE FOR ANY CLAIM,
	 DAMAGES OR OTHER LIABILITY, WHETHER IN AN ACTION OF CONTRACT, TORT OR OTHERWISE,
	 ARISING FROM, OUT OF OR IN CONNECTION WITH THE SOFTWARE OR THE USE OR OTHER
	 DEALINGS IN THE SOFTWARE.

	 You can contact me at:  terryeppler@gmail.com or eppler.terry@epa.gov

  </copyright>
  <summary>
	processing.py
  </summary>
  ******************************************************************************************
  '''
from __future__ import annotations
import string

import docx
from pymupdf import Page, Document
from sklearn.feature_extraction.text import TfidfVectorizer

from boogr import Error, ErrorDialog
from bs4 import BeautifulSoup
import chromadb
from chromadb.config import Settings
from collections import Counter
import html
import fitz
import glob
from gensim.models import Word2Vec
import json
import matplotlib.pyplot as plt
import numpy as np
import nltk
from nltk import FreqDist, ConditionalFreqDist
from nltk.corpus import stopwords, wordnet, words
from nltk.stem import WordNetLemmatizer, PorterStemmer
from nltk.tokenize import word_tokenize, sent_tokenize
import os
import pandas as pd
from pandas import DataFrame, Series
from pathlib import Path
from pinecone import Pinecone, ServerlessSpec
import re
import spacy
from spacy import Language
import sqlite3
from sentence_transformers import SentenceTransformer
from sklearn.decomposition import PCA
from sklearn.metrics.pairwise import cosine_similarity
from textblob import TextBlob
from typing import List, Optional, Dict, Tuple, Any, Set
import tiktoken
from tiktoken.core import Encoding
import unicodedata
from lxml import etree

try:
	nltk.data.find( 'tokenizers/punkt' )
except LookupError:
	nltk.download( 'punkt' )
	nltk.download( 'punkt_tab' )
	nltk.download( 'stopwords' )
	nltk.download( 'wordnet' )
	nltk.download( 'omw-1.4' )
	nltk.download( 'words' )

def throw_if( name: str, value: object ):
	if value is None:
		raise Exception( f'Argument "{name}" cannot be empty!' )

class Processor( ):
	'''
		
		Purpose:
		Base class for processing classes
		
	'''
	lemmatizer: Optional[ WordNetLemmatizer ]
	stemmer: Optional[ PorterStemmer ]
	file_path: Optional[ str ]
	normalized: Optional[ str ]
	lemmatized: Optional[ str ]
	tokenized: Optional[ str ]
	encoding: Optional[ Encoding ]
	nlp: Optional[ Language ]
	parts_of_speech: Optional[ List[ Tuple[ str, str ] ] ]
	embedddings: Optional[ List[ np.ndarray ] ]
	chunk_size: Optional[ int ]
	corrected: Optional[ str ]
	raw_input: Optional[ str ]
	raw_html: Optional[ str ]
	raw_pages: Optional[ List[ str ] ]
	lines: Optional[ List[ str ] ]
	tokens: Optional[ List[ str ] ]
	lines: Optional[ List[ str ] ]
	files: Optional[ List[ str ] ]
	pages: Optional[ List[ str ] ]
	paragraphs: Optional[ List[ str ] ]
	ids: Optional[ List[ int ] ]
	stop_words: Optional[ set ]
	vocabulary: Optional[ set ]
	corpus: Optional[ DataFrame ]
	removed: Optional[ List[ str ] ]
	frequency_distribution: Optional[ DataFrame ]
	
	def __init__( self ):
		self.lemmatizer = WordNetLemmatizer( )
		self.stemmer = PorterStemmer( )
		self.files = [ ]
		self.lines = [ ]
		self.tokens = [ ]
		self.lines = [ ]
		self.pages = [ ]
		self.ids = [ ]
		self.chunks = [ ]
		self.chunk_size = 0
		self.paragraphs = [ ]
		self.embedddings = [ ]
		self.stop_words = set( )
		self.vocabulary = set( )
		self.frequency_distribution = { }
		self.encoding = None
		self.corrected = None
		self.lowercase = None
		self.raw_html = None
		self.corpus = None
		self.file_path = ''
		self.raw_input = ''
		self.normalized = ''
		self.lemmatized = ''
		self.tokenized = ''
		self.cleaned_text = ''

class TextParser( Processor ):
	'''

		Purpose:
		---------
		Class providing path preprocessing functionality

		Methods:
		--------
		split_lines( self, path: str ) -> list
		split_pages( self, path: str, delimit: str ) -> list
		collapse_whitespace( self, path: str ) -> str
		remove_punctuation( self, path: str ) -> str:
		remove_special( self, path: str, keep_spaces: bool ) -> str:
		remove_html( self, path: str ) -> str
		remove_errors( self, path: str ) -> str
		remove_markdown( self, path: str ) -> str
		remove_stopwords( self, path: str ) -> str
		remove_headers( self, pages, min: int=3 ) -> str
		normalize_text( path: str ) -> str
		lemmatize_tokens( words: List[ str ] ) -> str
		tokenize_text( path: str ) -> str
		tokenize_words( path: str ) -> List[ str ]
		tokenize_sentences( path: str ) -> str
		chunk_text( self, path: str, max: int=800 ) -> List[ str ]
		chunk_words( self, path: str, max: int=800, over: int=50 ) -> List[ str ]
		split_paragraphs( self, path: str ) -> List[ str ]
		compute_frequency_distribution( self, words: List[ str ], proc: bool=True ) -> List[ str ]
		compute_conditional_distribution( self, words: List[ str ], condition: str=None,
		proc: bool=True ) -> List[ str ]
		create_vocabulary( self, frequency, min: int=1 ) -> List[ str ]
		create_wordbag( words: List[ str ] ) -> dict
		create_word2vec( sentences: List[ str ], vector_size=100, window=5, min_count=1 ) ->
		Word2Vec
		create_tfidf( words: List[ str ], max_features=1000, prep=True ) -> tuple

	'''
	lemmatizer: Optional[ WordNetLemmatizer ]
	stemmer: Optional[ PorterStemmer ]
	file_path: Optional[ str ]
	lowercase: Optional[ str ]
	normalized: Optional[ str ]
	lemmatized: Optional[ str ]
	tokenized: Optional[ str ]
	encoding: Optional[ Encoding ]
	nlp: Optional[ Language ]
	parts_of_speech: Optional[ List[ Tuple[ str, str ] ] ]
	embedddings: Optional[ List[ np.ndarray ] ]
	chunk_size: Optional[ int ]
	corrected: Optional[ str ]
	raw_input: Optional[ str ]
	raw_html: Optional[ str ]
	raw_pages: Optional[ List[ str ] ]
	lines: Optional[ List[ str ] ]
	tokens: Optional[ List[ str ] ]
	lines: Optional[ List[ str ] ]
	files: Optional[ List[ str ] ]
	pages: Optional[ List[ str ] ]
	paragraphs: Optional[ List[ str ] ]
	ids: Optional[ List[ int ] ]
	cleaned_text: Optional[ str ]
	cleaned_lines: Optional[ List[ str ] ]
	cleaned_tokens: Optional[ List[ str ] ]
	cleaned_pages: Optional[ List[ str ] ]
	cleaned_html: Optional[ str ]
	stop_words: Optional[ set ]
	vocabulary: Optional[ Series ]
	corpus: Optional[ DataFrame ]
	frequency_distribution: Optional[ DataFrame ]
	conditional_distribution: Optional[ DataFrame ]
	PUNCTUATION: Optional[ Set[ str ] ]
	CONTROL_CHARACTERS: Optional[ Set[ str ] ]

	def __init__( self ):
		'''

			Purpose:
			---------
			Constructor for 'Text' objects

		'''
		super( ).__init__( )
		self.PUNCTUATION = set( string.punctuation )
		self.CONTROL_CHARACTERS = ( {chr(i) for i in range(0x00, 0x20)} | {chr(0x7F)} )
		self.lemmatizer = WordNetLemmatizer( )
		self.stemmer = PorterStemmer( )
		self.encoding = tiktoken.get_encoding( 'cl100k_base' )
		self.lines = [ ]
		self.tokens = [ ]
		self.pages = [ ]
		self.ids = [ ]
		self.paragraphs = [ ]
		self.chunks = [ ]
		self.chunk_size = 10
		self.raw_pages = [ ]
		self.stop_words = set( )
		self.frequency_distribution = { }
		self.file_path = ''
		self.raw_input = ''
		self.normalized = ''
		self.lemmatized = ''
		self.tokenized = ''
		self.cleaned_text = ''
		self.vocabulary = None
		self.corrected = None
		self.lowercase = None
		self.raw_html = None
		self.translator = None
		self.tokenizer = None
		self.vectorizer = None
	
	def __dir__( self ) -> List[ str ] | None:
		'''

			Purpose:
			---------
			Provides a list of strings representing class members.


			Parameters:
			-----------
			- self

			Returns:
			--------
			- List[ str ] | None

		'''
		return [ # Attributes
				'file_path',
				'raw_input',
				'raw_pages',
				'normalized',
				'lemmatized',
				'tokenized',
				'corrected',
				'cleaned_text',
				'words',
				'paragraphs',
				'words',
				'pages',
				'chunks',
				'chunk_size',
				'stop_words',
				'removed',
				'lowercase',
				'encoding',
				'vocabulary',
				'translator',
				'lemmatizer',
				'stemmer',
				'tokenizer',
				'vectorizer',
				'conditional_distribution',
				# Methods
				'split_sentences',
				'split_pages',
				'collapse_whitespace',
				'compress_whitespace',
				'remove_punctuation',
				'remove_numbers',
				'remove_special',
				'remove_html',
				'remove_markdown',
				'remove_stopwords',
				'remove_formatting',
				'remove_headers',
				'remove_encodings',
				'tiktokenize',
				'lemmatize_text',
				'normalize_text',
				'chunk_text',
				'chunk_sentences',
				'chunk_files',
				'chunk_data',
				'chunk_datasets',
				'create_wordbag',
				'clean_file',
				'clean_files',
				'convert_jsonl',
				'speech_tagging',
				'split_paragraphs',
				'calculate_frequency_distribution',
				'create_vocabulary',
				'create_wordbag',
				'create_vectors',
				'encode_sentences',
				'semantic_search' ]
	
	def load_text( self, filepath: str ) -> str | None:
		"""
	`
			Purpose:
				Loads raw text from a file.
				
			Parameters:
				file_path (str): Path to the .txt file.
				
			Returns:
				str: Raw file content as string.`
			
		"""
		try:
			throw_if( 'filepath', filepath )
			if not os.path.exists( filepath ):
				raise FileNotFoundError( f'File not found: {filepath}' )
			else:
				self.file_path = filepath
			raw_text = open( self.file_path, mode='r', encoding='utf-8', errors='ignore' ).read()
			return raw_text
		except Exception as e:
			exception = Error( e )
			exception.module = 'processing'
			exception.cause = 'Text'
			exception.method = 'load_text( self, file_path: str ) -> str'
			error = ErrorDialog( exception )
			error.show( )
	
	def collapse_whitespace( self, text: str ) -> str | None:
		"""

			Purpose:
			-----------
			Removes extra lines from the string 'text'.

			Parameters:
			-----------
			- text : str

			Returns:
			--------
			A string with:
				- Consecutive whitespace reduced to a single space
				- Leading/trailing spaces removed
				- Blank words removed

		"""
		try:
			throw_if( 'text', text )
			self.raw_input = text
			extra_lines = re.sub( r'[\r\n]+', ' ', self.raw_input )
			self.cleaned_lines = [ line for line in extra_lines ]
			return ''.join( self.cleaned_lines )
		except Exception as e:
			exception = Error( e )
			exception.module = 'processing'
			exception.cause = 'Text'
			exception.method = 'collapse_whitespace( self, path: str ) -> str:'
			error = ErrorDialog( exception )
			error.show( )
	
	def compress_whitespace( self, text: str ) -> str | None:
		"""

			Purpose:
			-----------
			Removes extra spaces and blank words from the string 'text'.

			Parameters:
			-----------
			- text : str

			Returns:
			--------
			A string with:
				- Consecutive whitespace reduced to a single space
				- Leading/trailing spaces removed
				- Blank words removed

		"""
		try:
			throw_if( 'text', text )
			self.raw_input = text
			lines = [ ]
			for line in self.raw_input.splitlines( ):
				lines.append( re.sub( r"[ \t\s{2,1}]+", " ", line ).strip( ) )
			
			self.parsed_text = "\n".join( lines )
			return self.parsed_text
		except Exception as e:
			exception = Error( e )
			exception.module = 'processing'
			exception.cause = 'Text'
			exception.method = 'collapse_whitespace( self, path: str ) -> str:'
			error = ErrorDialog( exception )
			error.show( )
	
	def remove_punctuation( self, text: str ) -> str:
		"""
		
			Purpose:
			--------
			Removes all punctuation characters from the input text using NLTK's tokenizer.
		
			This function tokenizes the input into words and filters out any tokens that
			are composed entirely of punctuation characters.
		
			Parameters
			----------
			text : str
				The raw input text from which to remove punctuation.
		
			Returns
			-------
			str
				A cleaned string with all punctuation removed and original word spacing preserved.
		
			Example
			-------
			remove_punctuation("Hello, world! How's it going?")
			'Hello world Hows it going'
			
		"""
		try:
			throw_if( 'text', text )
			# Define sentence delimiters to preserve
			sentence_delimiters = {
					".": "__PERIOD__",
					"!": "__EXCLAMATION__",
					"?": "__QUESTION__",
					";": "__SEMICOLON__"
			}
			
			protected_text = text
			for delimiter, token in sentence_delimiters.items( ):
				protected_text = protected_text.replace( delimiter, token )
			
			# Remove all punctuation (Unicode-aware)
			without_punctuation = re.sub( r"[^\w\s]", " ", protected_text, flags=re.UNICODE )
			restored_text = without_punctuation
			for delimiter, token in sentence_delimiters.items( ):
				restored_text = restored_text.replace( token, delimiter )
			
			return restored_text
		except Exception as e:
			exception = Error( e )
			exception.module = 'processing'
			exception.cause = 'Text'
			exception.method = 'remove_punctuation( self, text: str ) -> str:'
			error = ErrorDialog( exception )
			error.show( )
			
	def normalize_text( self, text: str ) -> str | None:
		"""

			Purpose:
			-----------
			Converts the input 'text' to lower case

			This function:
			  - Converts the input 'text' to lower case

			Parameters:
			-----------
			- text : str
				The raw text

			Returns:
			--------
			- str
				A cleaned_lines path containing
				only letters, numbers, and spaces.

		"""
		try:
			throw_if( 'text', text )
			lower_cased = [ ]
			tokens = text.split( )
			for char in tokens:
				lower = char.lower( )
				lower_cased.append( lower )
			return ' '.join( lower_cased )
		except Exception as e:
			exception = Error( e )
			exception.module = 'chonky'
			exception.cause = 'Text'
			exception.method = 'normalize_text( self, text: str ) -> str:'
			error = ErrorDialog( exception )
			error.show( )
	
	def remove_errors( self, text: str  ) -> str:
		"""
		
			Purpose:
			----------
			Removes tokens that are not recognized as valid English words
			using the NLTK `words` corpus as a reference dictionary.
	
			This function is useful for cleaning text from OCR output, web-scraped data,
			or noisy documents by removing pseudo-words, typos, and out-of-vocabulary items.
	
			Parameters
			----------
			text : str
			The raw input text
	
			Returns
			-------
			str
			The raw input text without errors
	
			
		"""
		try:
			throw_if( 'text', text )
			_vocab = words.words( 'en' )
			_tokens = text.split(  )
			_words = [ w for w in _tokens if w in _vocab ]
			_data = ' '.join( _words )
			return _data
		except Exception as e:
			exception = Error( e )
			exception.module = 'processing'
			exception.cause = 'Text'
			exception.method = 'remove_errors( self, text: str  ) -> str'
			error = ErrorDialog( exception )
			error.show( )
	
	def remove_fragments( self, text: str ) -> str | None:
		"""

			Purpose:
			-----------
			Removes strings less that 4 chars in length


			Parameters:
			-----------
			- pages : str
				The raw path path path potentially
				containing special characters.

			Returns:
			--------
			- str
				A cleaned_lines path containing
				only letters, numbers, and spaces.

		"""
		try:
			throw_if( 'text', text )
			_cleaned = [ ]
			_fragments = text.split( )
			for char in _fragments:
				if char.isalpha( ) and len( char) > 2:
					_cleaned.append( char )
			return ' '.join( _cleaned )
		except Exception as e:
			exception = Error( e )
			exception.module = 'processing'
			exception.cause = 'Text'
			exception.method = 'remove_special( self, text: str ) -> str:'
			error = ErrorDialog( exception )
			error.show( )
	
	def remove_symbols( self, text: str ) -> str | None:
		"""

			Purpose:
			-----------
			Removes special characters from the path path path.

			This function:
			  - Retains only alphanumeric characters and whitespace
			  - Removes symbols like @, #, $, %, &, etc.
			  - Preserves letters, numbers, and spaces

			Parameters:
			-----------
			- pages : str
				The raw path path path potentially
				containing special characters.

			Returns:
			--------
			- str
				A cleaned_lines path containing
				only letters, numbers, and spaces.

		"""
		try:
			throw_if( 'text', text )
			self.raw_input = text
			MEANING_CRITICAL = {
					"&",  # AT&T, R&D, P&G, M&A
					"-",
					"--",
					"/",  # and/or, input/output
					"_",  # snake_case, file_names
					"%",  # percentage semantics
			}
			
			REMOVE_SYMBOLS = {
					"@",
					"#",
					"$",
					"^",
					"*",
					"=",
					"|",
					"\\",
					"<",
					">",
					"~"
			}
			
			cleaned = [ ]
			for ch in self.raw_input:
				if ch in MEANING_CRITICAL:
					cleaned.append( ' ' )
					continue
				if ch in REMOVE_SYMBOLS:
					continue
				cat = unicodedata.category( ch )
				if cat.startswith( "S" ):  # Symbol categories: Sc, Sk, Sm, So
					if ch in MEANING_CRITICAL:
						cleaned.append( ch )
					continue
				
				cleaned.append( ch )
			
			self.parsed_text = "".join( cleaned )
			return self.parsed_text
		except Exception as e:
			exception = Error( e )
			exception.module = 'processing'
			exception.cause = 'Text'
			exception.method = 'remove_special( self, text: str ) -> str:'
			error = ErrorDialog( exception )
			error.show( )
			
	def remove_html( self, text: str ) -> str | None:
		"""

			Purpose:
			--------
			Removes HTML tags from the path.

			This function:
			  - Parses the path as HTML
			  - Extracts and returns only the visible content without tags

			Parameters:
			-----------
			- pages : str
				The path path containing HTML tags.

			Returns:
			--------
			- str
				A cleaned_lines path with all HTML tags removed.

		"""
		try:
			throw_if( 'text', text )
			self.raw_html = text
			cleaned_html = BeautifulSoup( self.raw_html, 'html.parser' ).get_text( )
			return cleaned_html
		except Exception as e:
			exception = Error( e )
			exception.module = 'processing'
			exception.cause = 'Text'
			exception.method = 'remove_html( self, text: str ) -> str'
			error = ErrorDialog( exception )
			error.show( )
	
	def remove_xml( self, text: str ) -> str:
		"""
		
			Purpose:
			--------
			Remove XML tags from a string while preserving inner text content using
			lxml for robust, fast, and XPath-capable parsing.
	
			This function safely parses XML fragments by wrapping them in a synthetic
			root element, then extracts all textual content (element text and tails)
			while discarding tags and attributes.
	
			Parameters:
			-----------
			text : str
				Input text containing XML markup.
	
			Returns:
			--------
			str
				Text with XML tags removed and inner text preserved.
	
			Raises:
			-------
			RuntimeError
				Raised when XML parsing fails.
			
		"""
		if text is None:
			raise ValueError( "text cannot be None" )
		
		try:
			wrapped_text = f"<root>{text}</root>"
			parser = etree.XMLParser( recover=True, remove_comments=True,
				remove_blank_text=False )
			
			root = etree.fromstring( wrapped_text.encode( "utf-8" ), parser )
			text_parts = [ ]
			for element in root.iter( ):
				if element.text:
					text_parts.append( element.text )
				if element.tail:
					text_parts.append( element.tail )
			
			return "".join( text_parts )
		except Exception as e:
			exception = Error( e )
			exception.module = 'processing'
			exception.cause = 'Text'
			exception.method = 'remove_xml( self, text: str ) -> str'
			error = ErrorDialog( exception )
			error.show( )
	
	def remove_markdown( self, text: str ) -> str | None:
		"""


			Purpose:
			-----------
			Removes Markdown syntax (e.g., *, #, [], etc.)

			Parameters:
			-----------
			- pages : str
				The formatted path pages.

			Returns:
			--------
			- str
				A cleaned_lines version of the pages with formatting removed.

		"""
		try:
			throw_if( 'text', text )
			self.raw_input = text
			_text = re.sub( r'\[.*?]\(.*?\)', ' ', text )
			_unmarked = re.sub( r'[`_*#~><-]', ' ', _text )
			self.cleaned_text = re.sub( r'!\[.*?]\(.*?\)', ' ', _unmarked )
			return self.cleaned_text
		except Exception as e:
			exception = Error( e )
			exception.module = 'processing'
			exception.cause = 'Text'
			exception.method = 'remove_markdown( self, path: str ) -> str'
			error = ErrorDialog( exception )
			error.show( )
	
	def remove_stopwords( self, text: str ) -> str | None:
		"""

			Purpose:
			-----------
			This function:
			  - Removes English stopwords from the path pages path.
			  - Tokenizes the path pages
			  - Removes common stopwords (e.g., "the", "is", "and", etc.)
			  - Returns the pages with only meaningful words


			Parameters:
			-----------
			- pages : str
				The text string.

			Returns:
			--------
			- str
				A text string without stopwords.

		"""
		try:
			throw_if( 'text', text )
			self.stop_words = set( stopwords.words( 'english' ) )
			_words = text.split( None )
			_tokens = [ t for t in _words ]
			cleaned_tokens = [ w for w in _tokens if w not in self.stop_words ]
			cleaned_text = ' '.join( cleaned_tokens )
			return cleaned_text
		except Exception as e:
			exception = Error( e )
			exception.module = 'processing'
			exception.cause = 'Text'
			exception.method = 'remove_stopwords( self, text: str ) -> str'
			error = ErrorDialog( exception )
			error.show( )
	
	def remove_encodings( self, text: str ) -> str | None:
		"""

			Purpose:
			---------
			Cleans text of encoding artifacts by resolving HTML entities,
			Unicode escape sequences, and over-encoded byte strings.

			Parameters
			----------
			text : str
			Input string potentially containing encoded characters.

			Returns
			-------
			str
			Cleaned Unicode-normalized text.

		"""
		try:
			throw_if( 'text', text )
			try:
				text = bytes( text, 'utf-8' ).decode( 'unicode_escape' )
			except UnicodeDecodeError:
				pass

			self.raw_input = text
			_html = html.unescape( self.raw_input )
			_norm = unicodedata.normalize( 'NFKC', _html )
			_chars = re.sub( r'[\x00-\x1F\x7F]', '', _norm )
			cleaned_text = _chars.strip( )
			return cleaned_text
		except Exception as e:
			exception = Error( e )
			exception.module = 'processing'
			exception.cause = 'Text'
			exception.method = 'remove_encodings( self, text: str ) -> str'
			error = ErrorDialog( exception )
			error.show( )
	
	def remove_headers( self, filepath: str, lines: int=50, headers: int=3, footers: int=3, ) -> str:
		"""
		
			Purpose:
			--------
			Remove repetitive headers and footers from a text document by identifying recurring
			line blocks at page boundaries and stripping the most frequent candidates.
	
			Parameters:
			-----------
			filepath (str):
			Path to the plain-text document to clean.
			lines_per_page (int, optional):
			Approximate number of lines per page. Defaults to 50. Tune for your source.
			header_lines (int, optional):
			Number of top lines per page to consider a header candidate. Defaults to 3.
			footer_lines (int, optional):
			Number of bottom lines per page to consider a footer candidate. Defaults to 3.
	
			Returns:
			--------
			str:
			Cleaned document text with repetitive headers and footers removed.
				
		"""
		try:
			throw_if( 'filepath', filepath )
			if lines < 6:
				raise ValueError( 'Argument \"lines_per_page\" should be at least 6.' )
			if headers < 0 or footers < 0:
				msg = 'Arguments \"header_lines\" and \"footer_lines\" must be non-negative.'
				raise ValueError( msg )
			
			with open( filepath, 'r', encoding='utf-8', errors='ignore' ) as fh:
				all_lines: List[ str ] = fh.readlines( )
			
			pages = [ all_lines[ i: i + lines ] for i in
			          range( 0, len( all_lines ), lines ) ]
			
			header_counts = { }
			footer_counts = { }
			for page in pages:
				n = len( page )
				if n == 0:
					continue
				
				if headers > 0 and n >= headers:
					hdr = tuple( page[ :headers ] )
					if hdr in header_counts:
						header_counts[ hdr ] += 1
					else:
						header_counts[ hdr ] = 1
				
				if footers > 0 and n >= footers:
					ftr = tuple( page[ -footers: ] )
					if ftr in footer_counts:
						footer_counts[ ftr ] += 1
					else:
						footer_counts[ ftr ] = 1
			
			common_header = ( )
			if header_counts:
				common_header = max( header_counts.items( ), key=lambda kv: kv[ 1 ] )[ 0 ]
			
			common_footer = ( )
			if footer_counts:
				common_footer = max( footer_counts.items( ), key=lambda kv: kv[ 1 ] )[ 0 ]
			
			cleaned_pages = [ ]
			for page in pages:
				lines = list( page )
				
				if common_header and len( lines ) >= len( common_header ):
					if tuple( lines[ : len( common_header ) ] ) == common_header:
						lines = lines[ len( common_header ): ]
				
				if common_footer and len( lines ) >= len( common_footer ):
					if tuple( lines[ -len( common_footer ): ] ) == common_footer:
						lines = lines[ : -len( common_footer ) ]
				
				cleaned_pages.append( ''.join( lines ) )
			return '\n'.join( cleaned_pages )
		except Exception as e:
			exception = Error( e )
			exception.module = 'processing'
			exception.cause = 'Text'
			exception.method = 'remove_headers( self, filepath: str ) -> str'
			error = ErrorDialog( exception )
			error.show( )

	def remove_numbers( self, text: str ) -> str | None:
		"""

			Purpose:
			---------
			Removes the numbers 0 through 9 from the input text.

			Parameters
			----------
			text : str
			Input string potentially containing encoded characters.

			Returns
			-------
			str
			Cleaned Unicode-normalized text.

		"""
		try:
			throw_if( 'text', text )
			self.raw_input = text
			_clean = [ c for c in self.raw_input if not c.isdigit( ) ]
			self.cleaned_text = ''.join( _clean )
			return self.cleaned_text
		except Exception as e:
			exception = Error( e )
			exception.module = 'processing'
			exception.cause = 'Text'
			exception.method = 'remove_encodings( self, text: str ) -> str'
			error = ErrorDialog( exception )
			error.show( )
	
	def remove_numerals( self, text: str ) -> str | None:
		"""

			Purpose:
			---------
			Removes the numbers 0 through 9 from the input text.

			Parameters
			----------
			text : str
			Input string potentially containing encoded characters.

			Returns
			-------
			str
			Cleaned Unicode-normalized text.

		"""
		try:
			throw_if( "text", text )
			self.raw_input = text
			roman_pattern = (r"\bM{0,4}(CM|CD|D?C{0,3})"
			                 r"(XC|XL|L?X{0,3})"
			                 r"(IX|IV|V?I{0,3})\b" )
		
			self.parsed_text = re.sub( roman_pattern, " ", self.raw_input, flags=re.IGNORECASE, )
			return self.parsed_text
		except Exception as e:
			exception = Error( e )
			exception.module = 'processing'
			exception.cause = 'TextParser'
			exception.method = 'remove_numerals( self, text: str ) -> str'
			error = ErrorDialog( exception )
			error.show( )
	
	def remove_images( self, text: str ) -> str:
		"""
			Purpose:
			--------
			Remove image references from text, including Markdown images, HTML <img> tags,
			and standalone image URLs.
	
			Parameters:
			-----------
			text : str
				Input text.
	
			Returns:
			--------
			str
				Text with image references removed.
	
			Raises:
			-------
			RuntimeError
				Raised when processing fails.
		"""
		throw_if( "text", text )
		
		try:
			self.raw_input = text
			
			# Remove Markdown images: ![alt](path)
			without_markdown_images = re.sub(
				r"!\[[^\]]*]\([^)]*\)",
				" ",
				self.raw_input
			)
			
			# Remove HTML <img> tags
			without_html_images = re.sub(
				r"<img\b[^>]*>",
				" ",
				without_markdown_images,
				flags=re.IGNORECASE
			)
			
			# Remove standalone image URLs
			self.parsed_text = re.sub(
				r"https?://\S+\.(png|jpg|jpeg|gif|bmp|svg|webp)",
				" ",
				without_html_images,
				flags=re.IGNORECASE
			)
			
			return self.parsed_text
		except Exception as e:
			exception = Error( e )
			exception.module = 'processing'
			exception.cause = 'Text'
			exception.method = ('remove_formatting( self, text: str ) -> str')
			error = ErrorDialog( exception )
			error.show( )
	
	def lemmatize_text( self, text: str ) -> str | None:
		"""

			Purpose:
			--------
			Tokenizes input text into a list of word tokens.


			Parameters:
			-----------
			text (str): Input text to tokenize.


			Returns:
			--------
			List[str]: List of t strings.

		"""
		try:
			throw_if( 'text', text )
			_tokens = word_tokenize( text )
			_cleaned = [ self.lemmatizer.lemmatize( t ) for t in _tokens ]
			_lemmmatized = ' '.join( _cleaned )
			return _lemmmatized
		except Exception as e:
			exception = Error( e )
			exception.module = 'processing'
			exception.cause = 'Text'
			exception.method = 'lemmatize( self, text: str ) -> List[ str ]'
			error = ErrorDialog( exception )
			error.show( )
	
	def tiktokenize( self, text: str, encoding: str='cl100k_base' ) -> DataFrame:
		"""

			Purpose:
			---------
			Tokenizes text text into subword words using OpenAI's tiktoken tokenizer.
			This function leverages the tiktoken library, which provides byte-pair encoding (BPE)
			tokenization used in models such as GPT-3.5 and GPT-4. Unlike standard word
			tokenization,
			this function splits text into model-specific subword units.

			Parameters
			----------
			- text : str
				The text string to be tokenized.

			- model : str, optional
				The tokenizer model to use. Examples include 'cl100k_base' (default),
				'gpt-3.5-turbo', or 'gpt-4'. Ensure the model is supported by tiktoken.

			Returns
			-------
			- List[str]
				A list of string words representing BPE subword units.

		"""
		try:
			throw_if( 'text', text )
			self.encoding = tiktoken.get_encoding( encoding )
			token_ids = self.encoding.encode( text )
			_data = pd.DataFrame( token_ids )
			return _data
		except Exception as e:
			exception = Error( e )
			exception.module = 'processing'
			exception.cause = 'Text'
			exception.method = ('tiktokenize( self, text, encoding) -> List[ int ]')
			error = ErrorDialog( exception )
			error.show( )
	
	def speech_tagging( self, text: str ) -> List[ Tuple[ str, str ] ] | None:
		"""
		
			Purpose:
			--------
			Performs part-of-speech tagging.
		
		
			Parameters:
			-----------
			text (str): Text to process.
		
		
			Returns:
			--------
			List[Tuple[str, str]]: List of (token, POS tag) pairs.
			
		"""
		try:
			throw_if( 'text', text )
			self.nlp = spacy.load( 'en_core_web_sm' )
			self.parts_of_speech = [ ( token.text, token.pos_ ) for token in self.nlp( text ) ]
			return self.parts_of_speech
		except Exception as e:
			exception = Error( e )
			exception.module = 'processing'
			exception.cause = 'Text'
			exception.method = 'speech_tagging( self, text: str ) -> List[ Tuple[ str, str ] ]'
			error = ErrorDialog( exception )
			error.show( )
	
	def chunk_text( self, text: str, size: int=10 ) -> str:
		"""

			Purpose:
			-----------
			Tokenizes cleaned_lines pages and breaks it into chunks for downstream vectors.
			  - Converts pages to lowercase
			  - Tokenizes pages using NLTK's word_tokenize
			  - Breaks words into chunks of a specified size
			  - Optionally joins words into strings (for transformer models)

			Parameters:
			-----------
			- pages : str
				The cleaned_lines path pages to be tokenized and chunked.

			- size : int, optional (default=50)
				Number of words per chunk_words.

			- return_as_string : bool, optional (default=True)
				If True, returns each chunk_words as a path; otherwise, returns a get_list of
				words.

			Returns:
			--------
			- a list

		"""
		try:
			throw_if( 'text', text )
			_tokens = nltk.word_tokenize( text )
			_sentences = [ _tokens[ i: i + size ] for i in range( 0, len( _tokens ), size ) ]
			_datamap = [ ]
			for index, chunk in enumerate( _sentences ):
				_value =  ' '.join( chunk )
				_datamap.append( _value )
				
			_data = ' '.join( _datamap )
			return _data
		except Exception as e:
			exception = Error( e )
			exception.module = 'processing'
			exception.cause = 'Text'
			exception.method = 'chunk_sentences( self, text: str, max: int=10 ) -> DataFrame'
			error = ErrorDialog( exception )
			error.show( )
	
	def chunk_sentences( self, text: str, size: int=10 ) -> DataFrame:
		"""

			Purpose:
			-----------
			Tokenizes cleaned_lines pages and breaks it into chunks for downstream vectors.
			  - Converts pages to lowercase
			  - Tokenizes pages using NLTK's sent_tokenize
			  - Breaks words into chunks of a specified size
			  - Optionally joins words into strings (for transformer models)

			Parameters:
			-----------
			- pages : str
				The cleaned_lines path pages to be tokenized and chunked.

			- size : int, optional (default=50)
				Number of words per chunk_words.

			- return_as_string : bool, optional (default=True)
				If True, returns each chunk_words as a path; otherwise, returns a get_list of
				words.

			Returns:
			--------
			- a list

		"""
		try:
			throw_if( 'text', text )
			_tokens = nltk.word_tokenize( text )
			_sentences = [ _tokens[ i: i + size ] for i in range( 0, len( _tokens ), size ) ]
			_datamap = [ ]
			for i, c in enumerate( _sentences ):
				_item = ' '.join( c )
				_datamap.append( _item )
				
			_data = pd.DataFrame( _datamap )
			return _data
		except Exception as e:
			exception = Error( e )
			exception.module = 'processing'
			exception.cause = 'Text'
			exception.method = 'chunk_text( self, text: str, max: int=512 ) -> DataFrame'
			error = ErrorDialog( exception )
			error.show( )
	
	def split_sentences( self, text: str, size: int=10 ) -> List[ str ]:
		"""

			Purpose:
			________
			Splits the text string into a list of strings using NLTK's Punkt sentence tokenizer.
			This function is useful for preparing text for further processing,
			such as tokenization, parsing, or named entity recognition.

			Parameters
			----------
			- text : str
			The raw text string to be segmented into sentences.
			
			- size : int
			THe chunk size

			Returns
			-------
			- List of strings

		"""
		try:
			throw_if( 'text', text )
			_tokens = nltk.word_tokenize( text )
			_sentences = [ _tokens[ i: i + size ] for i in range( 0, len( _tokens ), size ) ]
			_data = [ ]
			for index, chunk in enumerate( _sentences ):
				_item = ' '.join( chunk )
				_data.append( _item )
				
			return _data
		except Exception as e:
			exception = Error( e )
			exception.module = 'processing'
			exception.cause = 'Text'
			exception.method = 'split_sentences( self, text: str ) -> DataFrame'
			error = ErrorDialog( exception )
			error.show( )
	
	def split_pages( self, filepath: str, num: int=50 ) -> List[ str ] | None:
		"""
		    
		    Purpose:
		    ---------
	        Splits a plain-text document into a list of pages using either form-feed characters
	        or fixed line count windows when form-feeds are absent.
		
		    Parameters:
		    -----------
	        file_path (str): Path to the text document.
	        lines_per_page (int): Approximate number of lines per page (default: 55).
		
		    Returns:
		    ---------
	        List[str]: List of page-level string segments.
		

		"""
		try:
			throw_if( 'filepath', filepath )
			if not os.path.exists( filepath ):
				raise FileNotFoundError( f'File not found: {filepath}' )
			else:
				self.file_path = filepath
			with open( self.file_path, 'r', encoding='utf-8', errors='ignore' ) as file:
				content = file.read( )
			if '\f' in content:
				return [ page.strip( ) for page in content.split( '\f' ) if page.strip( ) ]
			self.lines = content.splitlines( )
			i = 0
			n = len( self.lines )
			while i < n:
				page_lines = self.lines[ i: i + num ]
				page_text = '\n'.join( page_lines ).strip( )
				if page_text:
					self.pages.append( page_text )
				i += num
			return self.pages
		except Exception as e:
			exception = Error( e )
			exception.module = 'processing'
			exception.cause = 'Text'
			exception.method = 'split_pages( file_path )'
			error = ErrorDialog( exception )
			error.show( )
		
	def split_paragraphs( self, filepath: str ) -> DataFrame:
		"""

			Purpose:
			---------
			Reads  a file and splits it into paragraphs. A paragraph is defined as a block
			of path separated by one or more empty lines (eg, '\n\n').

			Parameters:
			-----------
			- path (str): Path to the path file.

			Returns:
			---------
			- list of str: List of paragraph strings.

		"""
		try:
			throw_if( 'filepath', filepath )
			if not os.path.exists( filepath ):
				raise FileNotFoundError( f'File not found: {filepath}' )
			else:
				self.file_path = filepath
			with open( self.file_path, 'r', encoding='utf-8', errors='ignore' ) as file:
				_input = file.read( )
				_paragraphs = [ pg.strip( ) for pg in _input.split( ' ' ) if pg.strip( ) ]
				_data = pd.DataFrame( _paragraphs )
				return _data
		except UnicodeDecodeError:
			with open( self.file_path, 'r', encoding='latin1', errors='ignore' ) as file:
				_input = file.read( )
				_paragraphs = [ pg.strip( ) for pg in _input.split( ' ' ) if pg.strip( ) ]
				_data = pd.DataFrame( _paragraphs )
				return _data
	
	def create_frequency_distribution( self, tokens: List[ str ] ) -> DataFrame:
		"""

			Purpose:
			--------
			Creates a word frequency freq_dist from a list of documents.

			Parameters:
			-----------
			- lines (list): List of raw or preprocessed path documents.
			- process (bool): Applies normalization, tokenization, stopword removal,
			and lemmatization.

			Returns:
			- dict: Dictionary of words and their corresponding frequencies.

		"""
		try:
			throw_if( 'tokens', tokens )
			self.tokens = tokens
			_freqdist = FreqDist( dict( Counter( self.tokens ) ) )
			_words = _freqdist.items( )
			_data = pd.DataFrame( _words, columns=[ 'Word', 'Frequency' ] )
			_data.index.name = 'ID'
			return _data
		except Exception as e:
			exception = Error( e )
			exception.module = 'processing'
			exception.cause = 'Text'
			exception.method = 'create_frequency_distribution(self, tokens: List[ str ])->DataFrame'
			error = ErrorDialog( exception )
			error.show( )
	
	def create_vocabulary( self, tokens: List[ str ], size: int=1 ) -> Series:
		"""

			Purpose:
			---------
			Builds a vocabulary list from a frequency
			distribution by applying a minimum frequency threshold.

			Parameters:
			-----------
			- freq_dist (dict):
				A dictionary mapping words to their frequencies.

			- min (int): Minimum num
				of occurrences required for a word to be included.

			Returns:
			--------
			- list: Sorted list of unique vocabulary words.

		"""
		try:
			throw_if( 'tokens', tokens )
			self.tokens = tokens
			_freqdist = FreqDist( dict( Counter( self.tokens ) ) )
			_vocab = _freqdist.items( )
			_vocabulary = pd.DataFrame( _vocab, columns=[ 'Word', 'Frequency' ] )
			_words = _vocabulary.iloc[ :, 0 ]
			return _words
		except Exception as e:
			exception = Error( e )
			exception.module = 'processing'
			exception.cause = 'Text'
			exception.method = ('create_vocabulary(self, freq_dist: dict, min: int=1)->List[str]')
			error = ErrorDialog( exception )
			error.show( )
	
	def create_wordbag( self, tokens: List[ str ] ) -> DataFrame:
		"""

			Purpose:
			--------
			Construct a Bag-of-Words (BoW) frequency dictionary from a list of strings.

			Parameters:
			-----------
			- words (list): List of words from a document.

			Returns:
			--------
			- dict: Word frequency dictionary.

		"""
		try:
			throw_if( 'tokens', tokens )
			self.tokens = tokens
			_freqdist = FreqDist( dict( Counter( self.tokens ) ) )
			_words = _freqdist.keys( )
			_data = pd.DataFrame( _words )
			return _data
		except Exception as e:
			exception = Error( e )
			exception.module = 'processing'
			exception.cause = 'Text'
			exception.method = 'create_wordbag( self, words: List[ str ] ) -> dict'
			error = ErrorDialog( exception )
			error.show( )
	
	def create_vectors( self, tokens: List[ str ] ) -> DataFrame:
		"""
		
			Purpose:
			----------
			Generates word embeddings using TF-IDF vectors from a list of tokens.
			This function treats each token as its own "document" to simulate a corpus
			and calculates a TF-IDF vector for each word based on term statistics.
		
			Parameters
			----------
			tokens : List[str]
				A list of individual word tokens to be embedded.
		
			Returns
			-------
			Dict[str, np.ndarray]
				A dictionary mapping each word to its corresponding TF-IDF vector.
		
			Example
			-------
			create_tfidf_word_embeddings(["nlp", "text", "nlp", "model"])
			{
				'nlp': array([0.7071, 0.0, 0.7071]),
				'text': array([0.0, 1.0, 0.0]),
				'model': array([0.0, 0.0, 1.0])
			}
		"""
		try:
			fake_docs = [ [ word ] for word in tokens ]
			joined_docs = [ ' '.join( doc ) for doc in fake_docs ]
			vectorizer = TfidfVectorizer( )
			X = vectorizer.fit_transform( joined_docs )
			feature_names = vectorizer.get_feature_names_out( )
			embeddings = { }
			for idx, word in enumerate( tokens ):
				vector = X[ idx ].toarray( ).flatten( )
				embeddings[ word ] = vector
			
			_data = pd.DataFrame( data=embeddings, columns=feature_names )
			return embeddings
		except Exception as e:
			exception = Error( e )
			exception.module = 'processing'
			exception.cause = 'Text'
			exception.method = 'create_vectors( self, tokens: List[str]) -> Dict[str, np.ndarray]'
			error = ErrorDialog( exception )
			error.show( )
			
	def clean_file( self, filepath: str ) -> str:
		"""

			Purpose:
			________
			Cleans text files given a source directory (src) and destination directory (dest)

			Parameters:
			----------
			- src (str): Source directory

			Returns:
			--------
			- None

		"""
		try:
			throw_if( 'filepath', filepath )
			if not os.path.exists( filepath ):
				raise FileNotFoundError( f'File not found: {filepath}' )
			else:
				_sourcepath = filepath
				_text = open( _sourcepath, 'r', encoding='utf-8', errors='ignore' ).read( )
				_collapsed = self.collapse_whitespace( _text )
				_compressed = self.compress_whitespace( _collapsed )
				_normalized = self.normalize_text( _compressed )
				_encoded = self.remove_encodings( _normalized )
				_special = self.remove_symbols( _encoded )
				_cleaned = self.remove_fragments( _special )
				_recompress = self.compress_whitespace( _cleaned )
				_lemmatized = self.lemmatize_text( _recompress )
				_stops = self.remove_stopwords( _lemmatized )
				return _stops
		except Exception as e:
			exception = Error( e )
			exception.module = 'processing'
			exception.cause = 'Text'
			exception.method = 'clean_file( self, src: str ) -> str'
			error = ErrorDialog( exception )
			error.show( )
	
	def clean_files( self, source: str, destination: str ) -> None:
		"""

			Purpose:
			________
			Cleans text files given a source directory (src) and destination directory (dest)

			Parameters:
			----------
			- src (str): Source directory
			- dest (str): Destination directory

			Returns:
			--------
			- None

		"""
		try:
			throw_if( 'src', source )
			throw_if( 'dest', destination )
			if not os.path.exists( source ):
				raise FileNotFoundError( f'File not found: {source}' )
			elif not os.path.exists( destination ):
				raise FileNotFoundError( f'File not found: {destination}' )
			else:
				_source = source
				_destpath = destination
				_files = os.listdir( _source )
				for f in _files:
					_processed = [ ]
					_filename = os.path.basename( f )
					_sourcepath = _source + '\\' + _filename
					_text = open( _sourcepath, 'r', encoding='utf-8', errors='ignore' ).read( )
					_collapsed = self.collapse_whitespace( _text )
					_compressed = self.compress_whitespace( _collapsed )
					_normalized = self.normalize_text( _compressed )
					_encoded = self.remove_encodings( _normalized )
					_special = self.remove_symbols( _encoded )
					_cleaned = self.remove_fragments( _special )
					_recompress = self.compress_whitespace( _cleaned )
					_lemmatized = self.lemmatize_text( _recompress )
					_stops = self.remove_stopwords( _lemmatized )
					_sentences = self.split_sentences( _stops )
					_destination = _destpath + '\\' + _filename
					_clean = open( _destination, 'wt', encoding='utf-8', errors='ignore' )
					_lines = ' '.join( _sentences )
					_clean.write( _lines )
		except Exception as e:
			exception = Error( e )
			exception.module = 'processing'
			exception.cause = 'Text'
			exception.method = 'clean_files( self, src: str, dest: str )'
			error = ErrorDialog( exception )
			error.show( )
	
	def chunk_files( self, source: str, destination: str ) -> None:
		"""

			Purpose:
			________
			Chunks cleaned text files given a _source directory and destination directory

			Parameters:
			----------
			- src (str): Source directory
			- dest (str): Destination directory

			Returns:
			--------
			- None

		"""
		try:
			throw_if( 'src', source )
			throw_if( 'dest', destination )
			if not os.path.exists( source ):
				raise FileNotFoundError( f'File not found: {source}' )
			elif not os.path.exists( destination ):
				raise FileNotFoundError( f'File not found: {destination}' )
			else:
				_source = source
				_destination = destination
				_files = os.listdir( _source )
				_words = [ ]
				for f in _files:
					_processed = [ ]
					_filename = os.path.basename( f )
					_sourcepath = _source + '\\' + _filename
					_text = open( _sourcepath, 'r', encoding='utf-8', errors='ignore' ).read( )
					_sentences =  self.split_sentences( _text )
					_datamap = [ ]
					for v in _sentences:
						_datamap.append( v )
						
					for s in _datamap:
						_processed.append( s )
					
					_final = _destination + '\\' + _filename
					_clean = open( _final, 'wt', encoding='utf-8', errors='ignore' )
					for p in _processed:
						_clean.write( p )
		except Exception as e:
			exception = Error( e )
			exception.module = 'processing'
			exception.cause = 'Text'
			exception.method = 'chunk_files( self, src: str, dest: str )'
			error = ErrorDialog( exception )
			error.show( )
	
	def chunk_data( self, filepath: str, size: int=10  ) -> DataFrame:
		"""

			Purpose:
			________
			Chunks cleaned text files given a source directory and destination directory

			Parameters:
			----------
			- src (str): Source directory

			Returns:
			--------
			- DataFrame

		"""
		try:
			throw_if( 'filepath', filepath )
			if not os.path.exists( filepath ):
				raise FileNotFoundError( f'File not found: {filepath}' )
			else:
				_source = filepath
				_processed = [ ]
				_wordlist = [ ]
				_vocab = words.words( 'en' )
				_text = open( _source, 'r', encoding='utf-8', errors='ignore' ).read( )
				_tokens = _text.split( )
				for s in _tokens:
					if s.isalpha( ) and s in _vocab:
						_wordlist.append( s )
				self.chunks = [ _wordlist[ i: i + size ] for i in range( 0, len( _wordlist ), size ) ]
				for i, c in enumerate( self.chunks ):
					_item =  '[' + ' '.join( c ) + '],'
					_processed.append( _item )
				_data = pd.DataFrame( _processed )
				return _data
		except Exception as e:
			exception = Error( e )
			exception.module = 'processing'
			exception.cause = 'Text'
			exception.method = 'chunk_data( self, filepath: str, size: int=512  ) -> DataFrame'
			error = ErrorDialog( exception )
			error.show( )
	
	def chunk_datasets( self, source: str, destination: str, size: int=10 ) -> DataFrame:
		"""

			Purpose:
			________
			Chunks cleaned text files given a source directory and destination directory

			Parameters:
			----------
			- src (str): Source directory

			Returns:
			--------
			- DataFrame

		"""
		try:
			throw_if( 'filepath', source )
			throw_if( 'destination', destination )
			if not os.path.exists( source ):
				raise FileNotFoundError( f'File not found: {source}' )
			elif not os.path.exists( destination ):
				raise FileNotFoundError( f'File not found: {destination}' )
			else:
				_src = source
				_destination = destination
				_files = os.listdir( _src )
				_words = [ ]
				for f in _files:
					_processed = [ ]
					_filename = os.path.basename( f )
					_sourcepath = _src+ '\\' + _filename
					_text = open( _sourcepath, 'r', encoding='utf-8', errors='ignore' ).read( )
					_collapsed = self.collapse_whitespace( _text )
					_compressed = self.compress_whitespace( _collapsed )
					_normalized = self.normalize_text( _compressed )
					_encoded = self.remove_encodings( _normalized )
					_special = self.remove_symbols( _encoded )
					_cleaned = self.remove_fragments( _special )
					_recompress = self.compress_whitespace( _cleaned )
					_lemmatized = self.lemmatize_text( _recompress )
					_stops = self.remove_stopwords( _lemmatized )
					_tokens = _stops.split( None )
					_chunks = [ _tokens[ i: i + size ] for i in range( 0, len( _tokens ), size ) ]
					_datamap = [ ]
					for i, c in enumerate( _chunks ):
						_row = ' '.join( c )
						_datamap.append( _row )
						
					for s in _datamap:
						_processed.append( s )
					
					_name = _filename.replace( '.txt', '.xlsx' )
					_savepath = ( _destination + f'\\' + _name )
					_data = pd.DataFrame( _processed, columns=[ 'Data', ] )
					_data.to_excel( _savepath, sheet_name='Dataset', index=False, columns=[ 'Data', ] )
		except Exception as e:
			exception = Error( e )
			exception.module = 'processing'
			exception.cause = 'Text'
			exception.method = 'chunk_data( self, filepath: str, size: int=15  ) -> DataFrame'
			error = ErrorDialog( exception )
			error.show( )

	def convert_jsonl( self, source: str, destination: str, size: int=10 ) -> None:
		"""

			Purpose:
			________
			Coverts text files to JSONL format given a source directory (Source)
			 and destination directory (destination)

			Parameters:
			--------
			- source (str): Source directory
			- destination (str): Destination directory

			Returns:
			--------
			- None

		"""
		try:
			throw_if( 'src', source )
			throw_if( 'dest', destination )
			if not os.path.exists( source ):
				raise FileNotFoundError( f'File not found: {source}' )
			elif not os.path.exists( destination ):
				raise FileNotFoundError( f'File not found: {destination}' )
			else:
				_source = source
				_destpath = destination
				_files = os.listdir( _source )
				_wordlist = [ ]
				for f in _files:
					_processed = [ ]
					_filename = os.path.basename( f )
					_sourcepath = _source + '\\' + _filename
					_text = open( _sourcepath, 'r', encoding='utf-8', errors='ignore' ).read( )
					_tokens =  _text.split( ' ' )
					_chunks = [ _tokens[ i: i + size ] for i in range( 0, len( _tokens ), size ) ]
					_datamap = [ ]
					for i, c in enumerate( _chunks ):
						_value = '{ ' +f' {i} : [ ' + ' '.join( c ) + ' ] }, ' + "\n"
						_datamap.append( _value )
						
					for s in _datamap:
						_processed.append( s )
					
					_destination = _destpath + '\\' + _filename.replace( '.txt', '.jsonl' )
					_clean = open( _destination, 'wt', encoding='utf-8', errors='ignore' )
					for p in _processed:
						_clean.write( p )
		except Exception as e:
			exception = Error( e )
			exception.module = 'processing'
			exception.cause = 'Text'
			exception.method = 'convert_jsonl( self, source: str, desination: str )'
			error = ErrorDialog( exception )
			error.show( )
	
	def encode_sentences( self, tokens: List[ str ], model: str='all-MiniLM-L6-v2' ) -> \
			Tuple[ List[ str ], np.ndarray ]:
		"""
		
			Purpose:
			Generate contextual sentence embeddings using SentenceTransformer.
			
			Parameters:
			sentences (list[str]): List of raw sentence strings.
			model_name (str): Model to use for encoding.
			
			Returns:
			tuple: (Cleaned sentences, Embeddings as np.ndarray)
			
		"""
		try:
			throw_if( 'tokens', tokens )
			throw_if( 'model', model )
			_transformer = SentenceTransformer( model )
			_tokens = [ self.lemmatizer.lemmatize( t ) for t in tokens ]
			_encoding = _transformer.encode( _tokens, show_progress_bar=True )
			return ( self.cleaned_tokens, np.array( _encoding ) )
		except Exception as e:
			exception = Error( e )
			exception.module = 'processing'
			exception.cause = 'Text'
			exception.method = 'encode_sentences( self, sentences: List[ str ], model_name ) -> ( )'
			error = ErrorDialog( exception )
			error.show( )
	
	def semantic_search( self, query: str, tokens: List[ str ], embeddings: np.ndarray,
			model: SentenceTransformer, top: int=5 ) -> List[ tuple[ str, float ] ]:
		"""
			Purpose:
				Perform semantic search over embedded corpus using query.
				
			Parameters:
				query (str): Natural language input.
				tokens (list[str]): Corpus sentences.
				embeddings (np.ndarray): Sentence embeddings.
				model (SentenceTransformer): Same model used for encoding.
				top (int): Number of matches to return.
				
			Returns:
				list[tuple[str, float]]: Top-k (sentence, similarity) pairs.
		"""
		try:
			throw_if( 'query', query )
			throw_if( 'tokens', tokens )
			throw_if( 'embedding', embeddings )
			throw_if( 'model', model )
			query_vec = model.encode( [ query ] )
			sims = cosine_similarity( query_vec, embeddings )[ 0 ]
			top_indices = sims.argsort( )[ ::-1 ][ : top ]
			return [ ( tokens[ i ], sims[ i ] ) for i in top_indices ]
		except Exception as e:
			exception = Error( e )
			exception.module = 'processing'
			exception.cause = 'Text'
			exception.method = ('semantic_search( self, query: str, tokens: List[ str ], '
			                    'embeddings: np.ndarray, model: SentenceTransformer,  '
			                    'top_k: int=5 ) -> List[ tuple[ str, float ] ]')
			error = ErrorDialog( exception )
			error.show( )
	

class WordParser( Processor ):
	"""

		Purpose:
		--------
		A class to extract, clean, and analyze text from Microsoft Word documents.

		Methods:
		--------
		split_sentences( self ) -> None
		clean_sentences( self ) -> None
		create_vocabulary( self ) -> None
		compute_frequency_distribution( self ) -> None
		summarize( self ) -> None:

	"""
	sentences: Optional[ List[ str ] ]
	cleaned_sentences: Optional[ List[ str ] ]
	document: Optional[ Document ]
	raw_text: Optional[ str ]
	paragraphs: Optional[ List[ str ] ]
	file_path: Optional[ str ]
	vocabulary: Optional[ set ]
	document: Optional[ Document ]
	
	def __init__( self, filepath: str ) -> None:
		"""

			Purpose:
			--------
			Initializes the WordTextProcessor with the path to the .docx file.

			Parameters:
			----------
			:param filepath: Path to the Microsoft Word document (.docx)

		"""
		super( ).__init__( )
		self.file_path = filepath
		self.raw_text = ''
		self.paragraphs = [ ]
		self.sentences = [ ]
		self.cleaned_sentences = [ ]
		self.vocabulary = set( )
	
	def __dir__( self ) -> List[ str ] | None:
		'''

			Purpose:
			---------
			Provides a list of strings representing class members.


			Parameters:
			-----------
			- self

			Returns:
			--------
			- List[ str ] | None

		'''
		return [ 'extract_text',
		         'split_sentences',
		         'clean_sentences',
		         'create_vocabulary',
		         'compute_frequency_distribution',
		         'summarize',
		         'filepath',
		         'raw_text',
		         'paragraphs',
		         'sentences',
		         'cleaned_sentences',
		         'vocabulary',
		         'freq_dist' ]
	
	def extract_text( self ) -> str | None:
		"""

			Purpose:
			--------
			Extracts raw text and paragraphs from the .docx file.

		"""
		try:
			self.document = Document( self.file_path )
			self.paragraphs = [ para.text.strip( ) for para in self.document.paragraphs if
				para.text.strip( ) ]
			self.raw_text = '\n'.join( self.paragraphs )
			return self.raw_text
		except Exception as e:
			exception = Error( e )
			exception.module = 'processing'
			exception.cause = 'Word'
			exception.method = 'extract_text( self ) -> str'
			error = ErrorDialog( exception )
			error.show( )
	
	def split_sentences( self ) -> List[ str ] | None:
		"""

			Purpose:
			--------
			Splits the raw text into sentences.

		"""
		try:
			self.sentences = sent_tokenize( self.raw_text )
			return self.sentences
		except Exception as e:
			exception = Error( e )
			exception.module = 'processing'
			exception.cause = 'Word'
			exception.method = 'split_sentences( self ) -> List[ str ]'
			error = ErrorDialog( exception )
			error.show( )
	
	def clean_sentences( self ) -> List[ str ] | None:
		"""

			Purpose:
			-------
			Cleans each _sentence: removes extra whitespace, punctuation, and lowers the text.

		"""
		try:
			for _sent in self.sentences:
				_sent = re.sub( r'[\r\n\t]+', ' ', _sent )
				_sent = re.sub( r"[^a-zA-Z0-9\s']", '', _sent )
				_sent = re.sub( r'\s{2,}', ' ', _sent ).strip( ).lower( )
				self.cleaned_sentences.append( _sent )
			return self.cleaned_sentences
		except Exception as e:
			exception = Error( e )
			exception.module = 'processing'
			exception.cause = 'Word'
			exception.method = 'clean_sentences( self ) -> List[ str ]'
			error = ErrorDialog( exception )
			error.show( )
	
	def create_vocabulary( self ) -> set | None:
		"""

			Purpose:
			--------
			Computes vocabulary terms from cleaned sentences.

		"""
		try:
			all_words = [ ]
			self.stop_words = set( stopwords.words( 'english' ) )
			for _sentence in self.cleaned_sentences:
				_tokens = word_tokenize( _sentence )
				self.tokens = [ token for token in _tokens if
					token.isalpha( ) and token not in self.stop_words ]
				all_words.extend( self.tokens )
			self.vocabulary = set( all_words )
			return self.vocabulary
		except Exception as e:
			exception = Error( e )
			exception.module = 'processing'
			exception.cause = 'Word'
			exception.method = 'create_vocabulary( self ) -> List[ str ]'
			error = ErrorDialog( exception )
			error.show( )
	
	def compute_frequency_distribution( self ) -> Dict[ str, int ] | None:
		"""

			Purpose:
			-------
			Computes frequency distribution of the vocabulary.

		"""
		try:
			words = [ ]
			for sentence in self.cleaned_sentences:
				tokens = word_tokenize( sentence )
				tokens = [ token for token in tokens if token.isalpha( ) ]
				words.extend( tokens )
			self.frequency_distribution = dict( Counter( words ) )
			return self.frequency_distribution
		except Exception as e:
			exception = Error( e )
			exception.module = 'processing'
			exception.cause = 'Word'
			exception.method = 'compute_frequency_distribution( self ) -> Dict[ str, int ]'
			error = ErrorDialog( exception )
			error.show( )
	
	def summarize( self ) -> List[ str ] | None:
		"""

			Purpose:
			-------
			Prints a summary of extracted and processed text.

		"""
		print( f'Document: {self.file_path}' )
		print( f'Paragraphs: {len( self.paragraphs )}' )
		print( f'Sentences: {len( self.sentences )}' )
		print( f'Vocabulary Size: {len( self.vocabulary )}' )
		print( f'Top 10 Frequent Words: {Counter( self.frequency_distribution ).most_common( 10 )}' )

class PdfParser( Processor ):
	"""

		Purpose:
		--------
		A utility class for extracting clean pages from PDF files into a list of strings.
		Handles nuances such as layout artifacts, page separation, optional filtering,
		and includes df detection capabilities.


		Methods:
		--------
		extract_lines( self, path, max: int=None) -> List[ str ]
		extract_text( self, path, max: int=None) -> str
		export_csv( self, tables: List[ pd.DataFrame ], filename: str=None ) -> None
		export_text( self, words: List[ str ], path: str=None ) -> None
		export_excel( self, tables: List[ pd.DataFrame ], path: str=None ) -> None

	"""
	strip_headers: Optional[ bool ]
	minimum_length: Optional[ int ]
	extract_tables: Optional[ bool ]
	extracted_lines: Optional[ List ]
	extracted_tables: Optional[ List ]
	extracted_pages: Optional[ List ]
	
	def __init__( self, headers: bool=False, size: int=10, tables: bool=True ) -> None:
		"""

			Purpose:
			-----------
			Initialize the PDF pages extractor with configurable settings.

			Parameters:
			-----------
			- headers (bool): If True, attempts to strip recurring headers/footers.
			- min (int): Minimum num of characters for a line to be included.
			- tables (bool): If True, extract pages from detected tables using block
			grouping.

		"""
		super( ).__init__( )
		self.strip_headers = headers
		self.minimum_length = size
		self.extract_tables = tables
		self.pages = [ ]
		self.lines = [ ]
		self.clean_lines = [ ]
		self.extracted_lines = [ ]
		self.extracted_tables = [ ]
		self.extracted_pages = [ ]
		self.tables = None
		self.file_path = ''
		self.page = ''
	
	def __dir__( self ) -> List[ str ] | None:
		'''

			Purpose:
			---------
			Provides a list of strings representing class members.


			Parameters:
			-----------
			- self

			Returns:
			--------
			- List[ str ] | None

		'''
		return [ 'strip_headers',
		         'minimum_length',
		         'extract_tables',
		         'file_path',
		         'page',
		         'pages',
		         'words',
		         'clean_lines',
		         'extracted_lines',
		         'extracted_tables',
		         'extracted_pages',
		         'extract_lines',
		         'extract_text',
		         'export_csv',
		         'export_text',
		         'export_excel' ]
	
	def extract_lines( self, path: str, size: Optional[ int ]=None ) -> List[ str ] | None:
		"""

			Purpose:
			--------
			Extract words of pages from a PDF,
			optionally limiting to the first N pages.

			Parameters:
			----------
			- path (str): Path to the PDF file
			- max (Optional[int]): Max num of pages to process (None for all pages)

			Returns:
			--------
			- List[str]: Cleaned list of non-empty words

		"""
		try:
			throw_if( 'path', path )
			self.file_path = path
			with fitz.open( self.file_path ) as doc:
				for i, page in enumerate( doc ):
					if size is not None and i >= size:
						break
					if self.extract_tables:
						page_lines = self._extract_tables( page )
					else:
						_text = page.get_text( 'text' )
						page_lines = _text.splitlines( )
					filtered = self._filter_lines( page_lines )
					self.extracted_lines.extend( filtered )
			return self.extracted_lines
		except Exception as e:
			exception = Error( e )
			exception.module = 'processing'
			exception.cause = 'PDF'
			exception.method = ('extract_lines( self, path: str, max: Optional[ int ]=None ) -> '
			                    'List[ str ]')
			error = ErrorDialog( exception )
			error.show( )
	
	def _extract_tables( self, page: Page ) -> List[ str ] | None:
		"""

			Purpose:
			--------
			Attempt to extract structured blocks
			such as tables using spatial grouping.

			Parameters:
			----------
			- page: PyMuPDF page object

			Returns:
			--------
			- List[str]: Grouped blocks including potential tables

		"""
		try:
			throw_if( 'page', page )
			tf = page.find_tables( )
			lines = [ ]
			for t in getattr( tf, 'tables', [ ] ):
				df = pd.DataFrame( t.extract( ) )  # or t.to_pandas()
				for row in df.values.tolist( ):
					lines.append( ' '.join( map( str, row ) ) )
			return lines
		except Exception as e:
			exception = Error( e )
			exception.module = 'processing'
			exception.cause = 'PDF'
			exception.method = '_extract_tables( self, page ) -> List[ str ]:'
			error = ErrorDialog( exception )
			error.show( )
	
	def _filter_lines( self, lines: List[ str ] ) -> List[ str ] | None:
		"""

			Purpose:
			-----------
			Filter and clean words from a page of pages.

			Parameters:
			- words (List[str]): Raw words of pages

			Returns:
			--------
			- List[str]: Filtered, non-trivial words

		"""
		try:
			if lines is None:
				raise Exception( 'The argument "lines" is None' )
			else:
				self.lines = lines
				clean = [ ]
				for line in lines:
					line = line.strip( )
					if len( line ) < self.minimum_length:
						continue
					if self.strip_headers and self._has_header( line ):
						continue
					clean.append( line )
				return clean
		except Exception as e:
			exception = Error( e )
			exception.module = 'processing'
			exception.cause = 'PDF'
			exception.method = '_filter_lines( self, words: List[ str ] ) -> List[ str ]'
			error = ErrorDialog( exception )
			error.show( )
	
	def _has_header( self, line: str ) -> bool | None:
		"""

			Purpose:
			--------
			Heuristic to detect common headers/footers (basic implementation).

			Parameters:
			- line (str): A line of pages

			Returns:
			--------
			- bool: True if line is likely a header or footer

		"""
		try:
			throw_if( 'line', line )
			_keywords = [ 'page', 'public law', 'u.s. government', 'united states' ]
			return any( kw in line.lower( ) for kw in _keywords )
		except Exception as e:
			exception = Error( e )
			exception.module = 'processing'
			exception.cause = 'PDF'
			exception.method = '_has_repeating_header( self, line: str ) -> bool'
			error = ErrorDialog( exception )
			error.show( )
	
	def extract_text( self, path: str, size: Optional[ int ]=None ) -> str | None:
		"""

			Purpose:
			---------
			Extract the entire pages from a PDF into one continuous path.

			Parameters:
			-----------
			- path (str): Path to the PDF file
			- max (Optional[int]): Maximum num of pages to process

			Returns:
			--------
			- str: Full concatenated pages

		"""
		try:
			throw_if( 'path', path )
			if size is not None and size > 0:
				self.file_path = path
				self.lines = self.extract_lines( self.file_path, size=size )
				return '\n'.join( self.lines )
			elif size is None or size <= 0:
				self.file_path = path
				self.lines = self.extract_lines( self.file_path )
				return '\n'.join( self.lines )
		except Exception as e:
			exception = Error( e )
			exception.module = 'processing'
			exception.cause = 'PDF'
			exception.method = 'extract_text( self, path: str, max: Optional[ int ]=None ) -> str:'
			error = ErrorDialog( exception )
			error.show( )
	
	def extract_tables( self, path: str, size: Optional[ int ]=None ) -> (
			List[ pd.DataFrame ] | None):
		"""

			Purpose:
			-----------
			Extract tables from the PDF and return them as a list of DataFrames.

			Parameters:
			- path (str): Path to the PDF file
			- max (Optional[int]): Maximum num of pages to process

			Returns:
			--------
			- List[pd.DataFrame]: List of DataFrames representing detected tables

		"""
		try:
			throw_if( 'path', path )
			throw_if( 'max', size )
			self.file_path = path
			self.tables = [ ]
			with fitz.open( self.file_path ) as doc:
				for i, page in enumerate( doc ):
					if size is not None and i >= size:
						break
					tf = page.find_tables( )
					for t in getattr( tf, 'tables', [ ] ):
						self.tables.append( pd.DataFrame( t.extract( ) ) )
			return self.tables
		except Exception as e:
			exception = Error( e )
			exception.module = 'processing'
			exception.cause = 'PDF'
			exception.method = (
				'extract_tables( self, path: str, max: Optional[ int ] = None ) -> List[ '
				'pd.DataFrame ]')
			error = ErrorDialog( exception )
			error.show( )
	
	def export_csv( self, tables: List[ pd.DataFrame ], filename: str ) -> None:
		"""

			Purpose:
			-----------
			Export a list of DataFrames (tables) to individual CSV files.

			Parameters:
			- tables (List[pd.DataFrame]): List of tables to export
			- filename (str): Prefix for output filenames (e.g., 'output_table')

		"""
		try:
			throw_if( 'tables', tables )
			throw_if( 'filename', filename )
			self.tables = tables
			for i, df in enumerate( self.tables ):
				df.to_csv( f'{filename}_{i + 1}.csv', index=False )
		except Exception as e:
			exception = Error( e )
			exception.module = 'processing'
			exception.cause = 'PDF'
			exception.method = (
				'export_csv( self, tables: List[ pd.DataFrame ], filename: str ) -> '
				'None')
			error = ErrorDialog( exception )
			error.show( )
	
	def export_text( self, lines: List[ str ], path: str ) -> None:
		"""

			Export extracted words of
			pages to a plain pages file.

			Parameters:
			-----------
			- words (List[str]): List of pages words
			- path (str): Path to output pages file

		"""
		try:
			throw_if( 'path', path )
			throw_if( 'lines', lines )
			self.file_path = path
			self.lines = lines
			with open( self.file_path, 'w', encoding='utf-8', errors='ignore' ) as f:
				for line in self.lines:
					f.write( line + '\n' )
		except Exception as e:
			exception = Error( e )
			exception.module = 'processing'
			exception.cause = 'PDF'
			exception.method = 'export_text( self, lines: List[ str ], path: str ) -> None'
			error = ErrorDialog( exception )
			error.show( )
	
	def export_excel( self, tables: List[ pd.DataFrame ], path: str ) -> None:
		"""

			Export all extracted tables into a single
			Excel workbook with one sheet per df.

			Parameters:
			-----------
			- tables (List[pd.DataFrame]): List of tables to export
			- path (str): Path to the output Excel file

		"""
		try:
			throw_if( 'tables', tables )
			self.tables = tables
			self.file_path = path
			with pd.ExcelWriter( self.file_path, engine='xlsxwriter' ) as _writer:
				for i, df in enumerate( self.tables ):
					_sheet = f'Table_{i + 1}'
					df.to_excel( _writer, sheet_name=_sheet, index=False )
		except Exception as e:
			exception = Error( e )
			exception.module = 'processing'
			exception.cause = 'PDF'
			exception.method = 'export_excel( self, tables: List[ pd.DataFrame ], path: str )->None'
			error = ErrorDialog( exception )
			error.show( )


