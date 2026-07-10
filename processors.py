'''
  ******************************************************************************************
	  Assembly:                Chonky
	  Filename:                processors.py
	  Author:                  Terry D. Eppler
	  Created:                 05-31-2022

	  Last Modified By:        Terry D. Eppler
	  Last Modified On:        05-01-2025
  ******************************************************************************************
  <copyright file="processors.py" company="Terry D. Eppler">

		 Boo is a df analysis tool that integrates various Generative AI, Text-Processing, and
		 Machine-Learning algorithms for federal analysts.
		 Copyright ©  2022  Terry Eppler

	 Permission is hereby granted, free of charge, to any person obtaining a copy
	 of this software and associated documentation files (the “Software”),
	 to deal in the Software without restriction,
	 including without limitation the rights to use,
	 copy, modify, merge, publish, distribute, sublicense,
	 and/or sell copies of the Software,
	 and to permit persons to whom the Software is furnished to do so,
	 subject to the following conditions:

	 The above copyright notice and this permission notice shall be included in all
	 copies or substantial portions of the Software.

	 THE SOFTWARE IS PROVIDED “AS IS”, WITHOUT WARRANTY OF ANY KIND, EXPRESS OR IMPLIED,
	 INCLUDING BUT NOT LIMITED TO THE WARRANTIES OF MERCHANTABILITY,
	 FITNESS FOR A PARTICULAR PURPOSE AND NON-INFRINGEMENT.
	 IN NO EVENT SHALL THE AUTHORS OR COPYRIGHT HOLDERS BE LIABLE FOR ANY CLAIM,
	 DAMAGES OR OTHER LIABILITY, WHETHER IN AN ACTION OF CONTRACT, TORT OR OTHERWISE,
	 ARISING FROM, OUT OF OR IN CONNECTION WITH THE SOFTWARE OR THE USE OR OTHER
	 DEALINGS IN THE SOFTWARE.

	 You can contact me at:  terryeppler@gmail.com or eppler.terry@epa.gov

  </copyright>
  <summary>
    Provides text-processing, parsing, tokenization, NLP, and PDF cleanup utilities for Chonky.

    Purpose:
        Defines the processing classes used after document ingestion and before embedding or
        vector persistence. The module supports text normalization, formatting cleanup,
        tokenization, lemmatization, stemming, vocabulary construction, frequency analysis,
        Word document parsing, and geometry-aware PDF text reconstruction. Existing wrapped
        exception handlers record validation and processing failures through the application
        logger before re-raising the project Error wrapper.
  </summary>
  ******************************************************************************************
  '''
from __future__ import annotations
import string

import docx
from pymupdf import Page, Document
from sklearn.feature_extraction.text import TfidfVectorizer
from boogr import Error, Logger
from bs4 import BeautifulSoup
import chromadb
from chromadb.config import Settings
from collections import Counter
import html
import fitz
import glob
from gensim.models import Word2Vec
import json
import matplotlib.pyplot as plt
import numpy as np
import nltk
from nltk import FreqDist, ConditionalFreqDist
from nltk.corpus import stopwords, wordnet, words
from nltk.stem import WordNetLemmatizer, PorterStemmer
from nltk.tokenize import word_tokenize, sent_tokenize
import os
import pandas as pd
from pandas import DataFrame, Series
from pathlib import Path
from pinecone import Pinecone, ServerlessSpec
import re
import spacy
from spacy import Language
import sqlite3
from sentence_transformers import SentenceTransformer
from sklearn.decomposition import PCA
from sklearn.metrics.pairwise import cosine_similarity
from textblob import TextBlob
from typing import List, Optional, Dict, Tuple, Any, Set
import tiktoken
from tiktoken.core import Encoding
import unicodedata
from lxml import etree

DELIMITERS: Set[ str ] = { '. ', '; ', '? ', '! ', ', ' }

SYMBOLS: Set[ str ] = { "@", "#", "$", "^", "*", "<", ">", "+", "=", "|", "\\", "<", ">", ":", "[",
	"]", "{", "}", "(", ")", "`", "~", "-", "_", '"', "'", ".", }

ASCII_LETTERS: Set[ str ] = set( string.ascii_letters )

DIGITS: Set[ str ] = set( string.digits )

PUNCTUATION: Set[ str ] = set( string.punctuation )

WHITESPACE: Set[ str ] = { " ", "\t", "\n", "\r", "\v", "\f" }

CONTROL_CHARACTERS: Set[ str ] = { chr( i ) for i in range( 0x00, 0x20 ) }.union( { chr( 0x7F ) } )

NUMERALS = (r"\bM{0,4}(CM|CD|D?C{0,3})"
            r"(XC|XL|L?X{0,3})"
            r"(IX|IV|V?I{0,3})\b")

try:
	nltk.data.find( 'tokenizers/punkt' )
except LookupError:
	nltk.download( 'punkt' )
	nltk.download( 'punkt_tab' )
	nltk.download( 'stopwords' )
	nltk.download( 'wordnet' )
	nltk.download( 'omw-1.4' )
	nltk.download( 'words' )

def throw_if( name: str, value: object ):
	if value is None:
		raise Exception( f'Argument "{name}" cannot be empty!' )

class Processor( ):
	"""Processor processing component.
	
	Purpose:
		Initializes shared state used by Chonky text, Word, NLTK, and PDF parser subclasses,
		including token caches, line/page buffers, vocabulary stores, NLP helpers,
		and cleaned-text fields used across the processing workflow.
	
	Attributes:
		lemmatizer: Runtime state used by ``Processor`` during Chonky processing operations.
		stemmer: Runtime state used by ``Processor`` during Chonky processing operations.
		file_path: Runtime state used by ``Processor`` during Chonky processing operations.
		normalized: Runtime state used by ``Processor`` during Chonky processing operations.
		lemmatized: Runtime state used by ``Processor`` during Chonky processing operations.
		tokenized: Runtime state used by ``Processor`` during Chonky processing operations.
		encoding: Runtime state used by ``Processor`` during Chonky processing operations.
		nlp: Runtime state used by ``Processor`` during Chonky processing operations.
		parts_of_speech: Runtime state used by ``Processor`` during Chonky processing operations.
		embedddings: Runtime state used by ``Processor`` during Chonky processing operations.
		chunk_size: Runtime state used by ``Processor`` during Chonky processing operations.
		corrected: Runtime state used by ``Processor`` during Chonky processing operations.
		raw_input: Runtime state used by ``Processor`` during Chonky processing operations.
		raw_html: Runtime state used by ``Processor`` during Chonky processing operations.
		raw_pages: Runtime state used by ``Processor`` during Chonky processing operations.
		lines: Runtime state used by ``Processor`` during Chonky processing operations.
		tokens: Runtime state used by ``Processor`` during Chonky processing operations.
		lines: Runtime state used by ``Processor`` during Chonky processing operations.
		files: Runtime state used by ``Processor`` during Chonky processing operations.
		pages: Runtime state used by ``Processor`` during Chonky processing operations.
		paragraphs: Runtime state used by ``Processor`` during Chonky processing operations.
		ids: Runtime state used by ``Processor`` during Chonky processing operations.
		stop_words: Runtime state used by ``Processor`` during Chonky processing operations.
		vocabulary: Runtime state used by ``Processor`` during Chonky processing operations.
		corpus: Runtime state used by ``Processor`` during Chonky processing operations.
		removed: Runtime state used by ``Processor`` during Chonky processing operations.
		frequency_distribution: Runtime state used by ``Processor`` during Chonky processing
		operations.
	"""
	lemmatizer: Optional[ WordNetLemmatizer ]
	stemmer: Optional[ PorterStemmer ]
	file_path: Optional[ str ]
	normalized: Optional[ str ]
	lemmatized: Optional[ str ]
	tokenized: Optional[ str ]
	encoding: Optional[ Encoding ]
	nlp: Optional[ Language ]
	parts_of_speech: Optional[ List[ Tuple[ str, str ] ] ]
	embedddings: Optional[ List[ np.ndarray ] ]
	chunk_size: Optional[ int ]
	corrected: Optional[ str ]
	raw_input: Optional[ str ]
	raw_html: Optional[ str ]
	raw_pages: Optional[ List[ str ] ]
	lines: Optional[ List[ str ] ]
	tokens: Optional[ List[ str ] ]
	lines: Optional[ List[ str ] ]
	files: Optional[ List[ str ] ]
	pages: Optional[ List[ str ] ]
	paragraphs: Optional[ List[ str ] ]
	ids: Optional[ List[ int ] ]
	stop_words: Optional[ set ]
	vocabulary: Optional[ set ]
	corpus: Optional[ DataFrame ]
	removed: Optional[ List[ str ] ]
	frequency_distribution: Optional[ DataFrame ]
	
	def __init__( self ):
		self.lemmatizer = WordNetLemmatizer( )
		self.stemmer = PorterStemmer( )
		self.files = [ ]
		self.lines = [ ]
		self.tokens = [ ]
		self.lines = [ ]
		self.pages = [ ]
		self.ids = [ ]
		self.chunks = [ ]
		self.chunk_size = 0
		self.paragraphs = [ ]
		self.embedddings = [ ]
		self.stop_words = set( )
		self.vocabulary = set( )
		self.frequency_distribution = { }
		self.encoding = None
		self.corrected = None
		self.lowercase = None
		self.raw_html = None
		self.corpus = None
		self.file_path = ''
		self.raw_input = ''
		self.normalized = ''
		self.lemmatized = ''
		self.tokenized = ''
		self.cleaned_text = ''

class TextParser( Processor ):
	"""TextParser processing component.
	
	Purpose:
		Provides text cleanup, normalization, tokenization, chunking, vocabulary,
		frequency-distribution, and semantic-preparation utilities used by the Text Processing,
		Analysis, and Tokenization tabs.
	
	Attributes:
		lowercase: Runtime state used by ``TextParser`` during Chonky processing operations.
		cleaned_text: Runtime state used by ``TextParser`` during Chonky processing operations.
		cleaned_lines: Runtime state used by ``TextParser`` during Chonky processing operations.
		cleaned_tokens: Runtime state used by ``TextParser`` during Chonky processing operations.
		cleaned_pages: Runtime state used by ``TextParser`` during Chonky processing operations.
		cleaned_html: Runtime state used by ``TextParser`` during Chonky processing operations.
		conditional_distribution: Runtime state used by ``TextParser`` during Chonky processing
		operations.
		PUNCTUATION: Runtime state used by ``TextParser`` during Chonky processing operations.
		CONTROL_CHARACTERS: Runtime state used by ``TextParser`` during Chonky processing
		operations.
		DELIMITERS: Runtime state used by ``TextParser`` during Chonky processing operations.
		DIGITS: Runtime state used by ``TextParser`` during Chonky processing operations.
		SYMBOLS: Runtime state used by ``TextParser`` during Chonky processing operations.
		NUMERALS: Runtime state used by ``TextParser`` during Chonky processing operations.
	"""
	lowercase: Optional[ str ]
	cleaned_text: Optional[ str ]
	cleaned_lines: Optional[ List[ str ] ]
	cleaned_tokens: Optional[ List[ str ] ]
	cleaned_pages: Optional[ List[ str ] ]
	cleaned_html: Optional[ str ]
	conditional_distribution: Optional[ DataFrame ]
	PUNCTUATION: Optional[ Set[ str ] ]
	CONTROL_CHARACTERS: Optional[ Set[ str ] ]
	DELIMITERS: Optional[ Set[ str ] ]
	DIGITS: Optional[ Set[ str ] ]
	SYMBOLS: Optional[ Set[ str ] ]
	NUMERALS: Optional[ str ]
	
	def __init__( self ):
		"""Initialize the TextParser instance.
		
		Purpose:
			Initializes parser state, reusable helper objects, and runtime caches used by later
			processing methods.
		"""
		super( ).__init__( )
		self.PUNCTUATION = PUNCTUATION
		self.CONTROL_CHARACTERS = ({ chr( i ) for i in range( 0x00, 0x20 ) } | { chr( 0x7F ) })
		self.DELIMITERS = DELIMITERS
		self.DIGITS = DIGITS
		self.SYMBOLS = SYMBOLS
		self.NUMERALS = NUMERALS
		self.lemmatizer = WordNetLemmatizer( )
		self.stemmer = PorterStemmer( )
		self.encoding = tiktoken.get_encoding( 'cl100k_base' )
		self.lines = [ ]
		self.tokens = [ ]
		self.pages = [ ]
		self.ids = [ ]
		self.paragraphs = [ ]
		self.chunks = [ ]
		self.chunk_size = 10
		self.raw_pages = [ ]
		self.stop_words = set( )
		self.frequency_distribution = { }
		self.file_path = ''
		self.raw_input = ''
		self.normalized = ''
		self.lemmatized = ''
		self.tokenized = ''
		self.cleaned_text = ''
		self.vocabulary = None
		self.corrected = None
		self.lowercase = None
		self.raw_html = None
		self.translator = None
		self.tokenizer = None
		self.vectorizer = None
	
	def __dir__( self ) -> List[ str ]:
		"""Return public parser members.
		
		Purpose:
			Returns the public attributes and methods exposed by the text parser for runtime
			introspection, diagnostics, and API documentation.
		
		Returns:
			List[ str ]: Public attribute and method names exposed by the parser.
		"""
		return [  # Attributes
			'PUNCTUATION', 'CONTROL_CHARACTERS', 'DELIMITERS', 'DIGITS', 'SYMBOLS', 'NUMERALS',
			'file_path', 'raw_input', 'raw_pages', 'normalized', 'lemmatized', 'tokenized',
			'corrected', 'cleaned_text', 'lines', 'tokens', 'pages', 'ids', 'paragraphs', 'chunks',
			'chunk_size', 'stop_words', 'frequency_distribution', 'lowercase', 'encoding',
			'vocabulary', 'raw_html', 'translator', 'lemmatizer', 'stemmer', 'tokenizer',
			'vectorizer',  # Methods
			'load_text', 'collapse_whitespace', 'remove_punctuation', 'reduce_repeats',
			'normalize_text', 'remove_errors', 'remove_fragments', 'remove_symbols', 'remove_html',
			'remove_xml', 'remove_markdown', 'remove_stopwords', 'remove_encodings',
			'remove_headers', 'remove_numbers', 'remove_numerals', 'remove_images', 'tiktokenize',
			'split_sentences', 'split_pages', 'split_paragraphs', 'create_frequency_distribution',
			'create_vocabulary', 'create_wordbag', 'create_vectors', 'clean_file', 'clean_files',
			'chunk_files', 'chunk_data', 'chunk_datasets', 'convert_jsonl', 'encode_sentences',
			'semantic_search' ]
	
	def load_text( self, filepath: str ) -> str | None:
		"""Load text.
		
		Purpose:
			Executes the ``load text`` operation for the ``TextParser`` workflow, updating
			instance state where required and returning the processed result used by downstream
			Chonky loading, processing, analysis, tokenization, or embedding steps.
		
		Args:
			filepath: Filepath value used by the processing operation. Expected type: ``str``.
		
		Returns:
			str | None: Result produced by the processing operation.
		
		Raises:
			Error: Raised when validation or processing fails.
		"""
		try:
			throw_if( 'filepath', filepath )
			if not os.path.exists( filepath ):
				raise FileNotFoundError( f'File not found: {filepath}' )
			else:
				self.file_path = filepath
			raw_text = open( self.file_path, mode='r', encoding='utf-8', errors='ignore' ).read( )
			return raw_text
		except Exception as e:
			exception = Error( e )
			exception.module = 'processors'
			exception.cause = 'TextParser'
			exception.method = 'load_text( self, file_path: str ) -> str'
			Logger( ).write( exception )
			raise exception
	
	def collapse_whitespace( self, text: str ) -> str | None:
		"""Collapse whitespace.
		
		Purpose:
			Converts each contiguous sequence of whitespace characters into one ordinary space
			while preserving capitalization, punctuation, numbers, symbols, and content order.
		
		Args:
			text: Source text containing whitespace to collapse.
		
		Returns:
			str | None: Text containing one ordinary space between adjacent content.
		
		Raises:
			Error: Raised when validation or whitespace collapsing fails.
		"""
		try:
			throw_if( 'text', text )
			return re.sub( r'\s+', ' ', text ).strip( )
		except Exception as e:
			exception = Error( e )
			exception.module = 'processors'
			exception.cause = 'TextParser'
			exception.method = 'collapse_whitespace( self, text: str ) -> str | None'
			Logger( ).write( exception )
			raise exception
	
	def remove_punctuation( self, text: str ) -> str:
		"""Remove punctuation.
		
		Purpose:
			Removes non-terminal punctuation while preserving sentence-ending periods, question
			marks, exclamation points, capitalization, numbers, and readable token boundaries.
		
		Args:
			text (str): Source text containing punctuation to remove.
		
		Returns:
			str: Text with non-terminal punctuation removed.
		
		Raises:
			Error: Raised when validation or punctuation removal fails.
		"""
		try:
			throw_if( 'text', text )
			_sentence_delimiters = { '.', '?', '!' }
			_characters: List[ str ] = [ ]
			
			for character in text:
				if character in _sentence_delimiters:
					_characters.append( character )
				elif character in self.PUNCTUATION:
					_characters.append( ' ' )
				else:
					_characters.append( character )
			
			_cleaned = ''.join( _characters )
			_cleaned = re.sub( r'[ \t]+', ' ', _cleaned )
			_cleaned = re.sub( r'\s+([.!?])', r'\1', _cleaned )
			_cleaned = re.sub( r'([.!?])(?=\w)', r'\1 ', _cleaned )
			return _cleaned.strip( )
		except Exception as e:
			exception = Error( e )
			exception.module = 'processors'
			exception.cause = 'TextParser'
			exception.method = 'remove_punctuation( self, text: str ) -> str'
			Logger( ).write( exception )
			raise exception
	
	def reduce_repeats( self, text: str ) -> str:
		"""Reduce repeated artifacts.
		
		Purpose:
			Reduces excessive repeated punctuation and symbol runs commonly produced by OCR or
			document extraction while preserving ellipses, paired dashes, capitalization, and text
			order.
		
		Args:
			text (str): Source text containing repeated punctuation or symbol artifacts.
		
		Returns:
			str: Text with excessive repeated artifacts reduced.
		
		Raises:
			Error: Raised when validation or repeat reduction fails.
		"""
		try:
			throw_if( 'text', text )
			_cleaned = re.sub( r'\.{4,}', '...', text )
			_cleaned = re.sub( r'-{3,}', '--', _cleaned )
			_cleaned = re.sub( r'([!?;,=:~@#$%^&*+|\\/])\1{2,}', r'\1', _cleaned )
			_cleaned = re.sub( r'[ \t]{2,}', ' ', _cleaned )
			return _cleaned
		except Exception as e:
			exception = Error( e )
			exception.module = 'processors'
			exception.cause = 'TextParser'
			exception.method = 'reduce_repeats( self, text: str ) -> str'
			Logger( ).write( exception )
			raise exception
	
	def normalize_text( self, text: str ) -> str | None:
		"""Normalize text case.
		
		Purpose:
			Converts alphabetic text to lowercase while preserving punctuation, numbers,
			whitespace,
			document structure, and all non-case-related content.
		
		Args:
			text (str): Source text to convert to lowercase.
		
		Returns:
			str | None: Lowercase representation of the supplied text.
		
		Raises:
			Error: Raised when validation or lowercase normalization fails.
		"""
		try:
			throw_if( 'text', text )
			return text.lower( )
		except Exception as e:
			exception = Error( e )
			exception.module = 'processors'
			exception.cause = 'TextParser'
			exception.method = 'normalize_text( self, text: str ) -> str'
			Logger( ).write( exception )
			raise exception
	
	def remove_errors( self, text: str ) -> str:
		"""Remove extraction errors.
		
		Purpose:
			Repairs recognized mojibake sequences and removes replacement characters, null bytes,
			invisible formatting artifacts, surrogate characters, private-use extraction glyphs,
			and unsupported control characters while preserving valid text, punctuation, numbers,
			line boundaries, tabs, and document order.
		
		Args:
			text: Source text containing extraction or encoding artifacts.
		
		Returns:
			str: Text with recognized extraction artifacts removed or repaired.
		
		Raises:
			Error: Raised when validation or extraction-error removal fails.
		"""
		try:
			throw_if( 'text', text )
			cleaned = text
			replacements = { '\ufffd': '', '\ufeff': '', '\u200b': '', '\u200c': '', '\u200d': '',
				'\u2060': '', '\x00': '', 'ï»¿': '', 'Â ': ' ', 'â€“': '–', 'â€”': '—', 'â€˜': '‘',
				'â€™': '’', 'â€œ': '“', 'â€\x9d': '”' }
			
			for invalid_value, replacement_value in replacements.items( ):
				cleaned = cleaned.replace( invalid_value, replacement_value )
			
			cleaned_characters: List[ str ] = [ ]
			for character in cleaned:
				if character in { '\n', '\r', '\t' }:
					cleaned_characters.append( character )
					continue
				
				category = unicodedata.category( character )
				if category in { 'Cc', 'Cf', 'Co', 'Cs' }:
					continue
				
				cleaned_characters.append( character )
			
			return ''.join( cleaned_characters )
		except Exception as e:
			exception = Error( e )
			exception.module = 'processors'
			exception.cause = 'TextParser'
			exception.method = 'remove_errors( self, text: str ) -> str'
			Logger( ).write( exception )
			raise exception
	
	def remove_fragments( self, text: str ) -> str:
		"""Remove malformed fragments.
		
		Purpose:
			Removes isolated extraction debris and unsupported private-use glyphs while preserving
			valid words, acronyms, identifiers, legal symbols, punctuation, numbers, whitespace,
			and document order.
		
		Args:
			text: Source text containing possible extraction fragments.
		
		Returns:
			str: Text with isolated malformed extraction fragments removed.
		
		Raises:
			Error: Raised when validation or fragment removal fails.
		"""
		try:
			throw_if( 'text', text )
			tokens = re.findall( r'\S+|\s+', text )
			cleaned_tokens: List[ str ] = [ ]
			preserved_fragments = { '§', '§§', '¶', '©', '®', '™', '%', '&', '+', '=', '$', '€',
				'£', '¥', '#' }
			
			extraction_fragments = { '•', '◦', '▪', '▫', '■', '□', '●', '○', '◆', '◇', '►', '▶',
				'➢', '➤', '', '', '�' }
			
			for token in tokens:
				if token.isspace( ):
					cleaned_tokens.append( token )
					continue
				
				visible = ''.join( character for character in token if
					character != '\ufffd' and unicodedata.category( character ) not in { 'Cf', 
						'Co',
						'Cs' } )
				
				if not visible:
					continue
				
				if visible in extraction_fragments:
					continue
				
				if any( character.isalnum( ) for character in visible ):
					cleaned_tokens.append( visible )
					continue
				
				if visible in preserved_fragments:
					cleaned_tokens.append( visible )
					continue
				
				if any( unicodedata.category( character ).startswith( 'P' ) for character in
						visible ):
					cleaned_tokens.append( visible )
			
			return ''.join( cleaned_tokens )
		except Exception as e:
			exception = Error( e )
			exception.module = 'processors'
			exception.cause = 'TextParser'
			exception.method = 'remove_fragments( self, text: str ) -> str'
			Logger( ).write( exception )
			raise exception
	
	def remove_symbols( self, text: str ) -> str | None:
		"""Remove configured symbols.
		
		Purpose:
			Removes configured non-sentence symbols while preserving capitalization, numbers,
			sentence delimiters, document order, and readable token boundaries.
		
		Args:
			text (str): Source text containing configured symbols to remove.
		
		Returns:
			str | None: Text with configured symbols removed.
		
		Raises:
			Error: Raised when validation or symbol removal fails.
		"""
		try:
			throw_if( 'text', text )
			_sentence_delimiters = { '.', '?', '!', ';', ':' }
			_symbols = self.SYMBOLS.difference( _sentence_delimiters )
			_characters: List[ str ] = [ ]
			for character in text:
				if character in _symbols:
					_characters.append( ' ' )
				else:
					_characters.append( character )
			
			_cleaned = ''.join( _characters )
			_cleaned = re.sub( r'[ \t]{2,}', ' ', _cleaned )
			return _cleaned.strip( )
		except Exception as e:
			exception = Error( e )
			exception.module = 'processors'
			exception.cause = 'TextParser'
			exception.method = 'remove_symbols( self, text: str ) -> str'
			Logger( ).write( exception )
			raise exception
	
	def remove_html( self, text: str ) -> str:
		"""Remove HTML markup.
		
		Purpose:
			Removes HTML elements and non-visible content while preserving visible text, inline
			word boundaries, paragraph boundaries, capitalization, punctuation, numbers, URLs,
			legal citations, and document order.
		
		Args:
			text: Source text containing HTML markup or entities.
		
		Returns:
			str: Visible text extracted from the supplied HTML content.
		
		Raises:
			Error: Raised when validation or HTML processing fails.
		"""
		try:
			throw_if( 'text', text )
			self.raw_html = text
			
			soup = BeautifulSoup( self.raw_html, 'html.parser' )
			
			for element in soup( [ 'script', 'style', 'noscript', 'template' ] ):
				element.decompose( )
			
			block_tags = { 'address', 'article', 'aside', 'blockquote', 'br', 'dd', 'div', 'dl',
				'dt', 'fieldset', 'figcaption', 'figure', 'footer', 'form', 'h1', 'h2', 'h3', 'h4',
				'h5', 'h6', 'header', 'hr', 'li', 'main', 'nav', 'ol', 'p', 'pre', 'section',
				'table', 'tbody', 'td', 'tfoot', 'th', 'thead', 'tr', 'ul' }
			
			for element in soup.find_all( block_tags ):
				if element.name in { 'br', 'hr' }:
					element.replace_with( '\n' )
				else:
					element.insert_before( '\n' )
					element.insert_after( '\n' )
			
			visible_text = soup.get_text( separator='' )
			visible_text = html.unescape( visible_text )
			visible_text = visible_text.replace( '\u00a0', ' ' )
			visible_text = re.sub( r'[ \t]+', ' ', visible_text )
			visible_text = re.sub( r' *\n *', '\n', visible_text )
			visible_text = re.sub( r'\n{3,}', '\n\n', visible_text )
			
			self.cleaned_text = visible_text.strip( )
			return self.cleaned_text
		except Exception as e:
			exception = Error( e )
			exception.module = 'processors'
			exception.cause = 'TextParser'
			exception.method = 'remove_html( self, text: str ) -> str'
			Logger( ).write( exception )
			raise exception
	
	def remove_xml( self, text: str ) -> str:
		"""Remove XML markup.
		
		Purpose:
			Removes XML declarations, comments, processing instructions, and elements while
			preserving visible text, inline word boundaries, existing line boundaries,
			capitalization, punctuation, numbers, URLs, legal citations, and document order.
		
		Args:
			text: Source text containing XML markup.
		
		Returns:
			str: Visible text extracted from the supplied XML content.
		
		Raises:
			Error: Raised when validation or XML processing fails.
		"""
		try:
			throw_if( 'text', text )
			
			source_text = re.sub( r'^\s*<\?xml[^>]*\?>', '', text, count=1, flags=re.IGNORECASE )
			
			wrapped_text = f'<root>{source_text}</root>'
			parser = etree.XMLParser( recover=True, remove_comments=True, remove_blank_text=False,
				resolve_entities=False, no_network=True )
			
			root = etree.fromstring( wrapped_text.encode( 'utf-8' ), parser )
			
			if root is None:
				return text
			
			text_parts: List[ str ] = [ ]
			
			for value in root.itertext( ):
				if not isinstance( value, str ) or not value:
					continue
				
				value = value.replace( '\r\n', '\n' )
				value = value.replace( '\r', '\n' )
				
				if (text_parts and text_parts[ -1 ] and not text_parts[ -1 ][
					-1 ].isspace( ) and not value[ 0 ].isspace( ) and text_parts[ -1 ][
					-1 ].isalnum( ) and value[ 0 ].isalnum( )):
					text_parts.append( ' ' )
				
				text_parts.append( value )
			
			cleaned_text = ''.join( text_parts )
			cleaned_text = re.sub( r'[ \t\f\v]+', ' ', cleaned_text )
			cleaned_text = re.sub( r' *\n *', '\n', cleaned_text )
			cleaned_text = re.sub( r'\n{3,}', '\n\n', cleaned_text )
			
			return cleaned_text.strip( )
		except Exception as e:
			exception = Error( e )
			exception.module = 'processors'
			exception.cause = 'TextParser'
			exception.method = 'remove_xml( self, text: str ) -> str'
			Logger( ).write( exception )
			raise exception
	
	def remove_markdown( self, text: str ) -> str | None:
		"""Remove Markdown syntax.
		
		Purpose:
			Removes Markdown formatting syntax while preserving visible link labels, image
			alternative text, headings, list content, code content, capitalization, hyphenated
			terms, and paragraph structure.
		
		Args:
			text (str): Source text containing Markdown syntax.
		
		Returns:
			str | None: Readable text with Markdown formatting syntax removed.
		
		Raises:
			Error: Raised when validation or Markdown processing fails.
		"""
		try:
			throw_if( 'text', text )
			self.raw_input = text
			_text = re.sub( r'!\[([^\]]*)\]\([^)]+\)', r'\1', self.raw_input )
			_text = re.sub( r'\[([^\]]+)\]\([^)]+\)', r'\1', _text )
			_text = re.sub( r'!\[([^\]]*)\]\[[^\]]*\]', r'\1', _text )
			_text = re.sub( r'\[([^\]]+)\]\[[^\]]*\]', r'\1', _text )
			_text = re.sub( r'^\s{0,3}#{1,6}\s+', '', _text, flags=re.MULTILINE )
			_text = re.sub( r'^\s{0,3}>\s?', '', _text, flags=re.MULTILINE )
			_text = re.sub( r'^\s*[-+*]\s+', '', _text, flags=re.MULTILINE )
			_text = re.sub( r'^\s*\d+[.)]\s+', '', _text, flags=re.MULTILINE )
			_text = re.sub( r'^\s*(```+|~~~+)[^\n]*$', '', _text, flags=re.MULTILINE )
			_text = re.sub( r'`([^`\n]+)`', r'\1', _text )
			_text = re.sub( r'(\*\*|__)(.*?)\1', r'\2', _text )
			_text = re.sub( r'(?<!\*)\*([^*\n]+)\*(?!\*)', r'\1', _text )
			_text = re.sub( r'(?<!_)_([^_\n]+)_(?!_)', r'\1', _text )
			_text = re.sub( r'~~(.*?)~~', r'\1', _text )
			_text = re.sub( r'^\s*([-*_])(?:\s*\1){2,}\s*$', '', _text, flags=re.MULTILINE )
			_text = re.sub( r'^\s*\[[^\]]+\]:\s+\S+.*$', '', _text, flags=re.MULTILINE )
			self.cleaned_text = _text
			return self.cleaned_text
		except Exception as e:
			exception = Error( e )
			exception.module = 'processors'
			exception.cause = 'TextParser'
			exception.method = 'remove_markdown( self, text: str ) -> str'
			Logger( ).write( exception )
			raise exception
	
	def remove_stopwords( self, text: str ) -> str:
		"""Remove English stopwords.
		
		Purpose:
			Removes standalone English stopwords while preserving URLs, punctuation, numbers,
			line boundaries, paragraph boundaries, capitalization, citations, and document order.
		
		Args:
			text: Source text containing English stopwords.
		
		Returns:
			str: Text with standalone English stopwords removed.
		
		Raises:
			Error: Raised when validation or stopword removal fails.
		"""
		try:
			throw_if( 'text', text )
			stop_words = set( stopwords.words( 'english' ) )
			url_pattern = re.compile( r'https?://[^\s<>"\']+|www\.[^\s<>"\']+',
				flags=re.IGNORECASE )
			word_pattern = re.compile( r"\b[A-Za-z]+(?:['’][A-Za-z]+)*\b" )
			text_parts: List[ str ] = [ ]
			start_index = 0
			for url_match in url_pattern.finditer( text ):
				segment = text[ start_index:url_match.start( ) ]
				
				segment = word_pattern.sub( lambda match: (
					'' if match.group( 0 ).lower( ) in stop_words else match.group( 0 )), segment )
				
				text_parts.append( segment )
				text_parts.append( url_match.group( 0 ) )
				start_index = url_match.end( )
			
			final_segment = text[ start_index: ]
			final_segment = word_pattern.sub(
				lambda match: ('' if match.group( 0 ).lower( ) in stop_words else match.group( 0
				)), final_segment )
			text_parts.append( final_segment )
			cleaned_text = ''.join( text_parts )
			cleaned_text = re.sub( r'[ \t]{2,}', ' ', cleaned_text )
			cleaned_text = re.sub( r'[ \t]+([,.;:!?%)\]}])', r'\1', cleaned_text )
			cleaned_text = re.sub( r'([({\[])[ \t]+', r'\1', cleaned_text )
			cleaned_text = re.sub( r'[ \t]+\n', '\n', cleaned_text )
			cleaned_text = re.sub( r'\n[ \t]+', '\n', cleaned_text )
			return cleaned_text.strip( )
		except Exception as e:
			exception = Error( e )
			exception.module = 'processors'
			exception.cause = 'TextParser'
			exception.method = 'remove_stopwords( self, text: str ) -> str'
			Logger( ).write( exception )
			raise exception
	
	def remove_encodings( self, text: str ) -> str:
		"""Remove encoding artifacts.
		
		Purpose:
			Decodes HTML entities, normalizes compatible Unicode characters, converts Unicode
			spacing characters to ordinary spaces, and removes control, formatting, surrogate,
			private-use, replacement, and soft-hyphen artifacts while preserving readable text,
			capitalization, punctuation, numbers, tabs, line boundaries, URLs, and document order.
		
		Args:
			text: Source text containing encoded entities or Unicode artifacts.
		
		Returns:
			str: Text containing normalized Unicode characters without unsafe encoding artifacts.
		
		Raises:
			Error: Raised when validation or encoding normalization fails.
		"""
		try:
			throw_if( 'text', text )
			self.raw_input = text
			decoded = html.unescape( self.raw_input )
			normalized = unicodedata.normalize( 'NFKC', decoded )
			spacing_characters = { '\u00a0', '\u1680', '\u180e', '\u2000', '\u2001', '\u2002',
				'\u2003', '\u2004', '\u2005', '\u2006', '\u2007', '\u2008', '\u2009', '\u200a',
				'\u200b', '\u202f', '\u205f', '\u3000' }
			
			removable_characters = { '\u00ad', '\u200c', '\u200d', '\u2060', '\ufeff', '\ufffd' }
			
			cleaned_characters: List[ str ] = [ ]
			for character in normalized:
				if character in { '\n', '\r', '\t' }:
					cleaned_characters.append( character )
					continue
				
				if character in spacing_characters:
					cleaned_characters.append( ' ' )
					continue
				
				if character in removable_characters:
					continue
				
				category = unicodedata.category( character )
				if category in { 'Cc', 'Cf', 'Co', 'Cs' }:
					continue
				
				cleaned_characters.append( character )
			
			return ''.join( cleaned_characters ).strip( )
		except Exception as e:
			exception = Error( e )
			exception.module = 'processors'
			exception.cause = 'TextParser'
			exception.method = 'remove_encodings( self, text: str ) -> str'
			Logger( ).write( exception )
			raise exception
	
	def remove_headers( self, filepath: str, lines: int=50, headers: int=3,
		footers: int=3 ) -> str:
		"""Remove repeated headers and footers.
		
		Purpose:
			Removes exact header and footer line sequences repeated across multiple fixed-length
			text pages while preserving unique page content, line boundaries, page order, and
			single-page documents.
		
		Args:
			filepath: Path to the UTF-8 text file to process.
			lines: Maximum number of lines assigned to each logical page.
			headers: Number of leading lines evaluated as a repeated header.
			footers: Number of trailing lines evaluated as a repeated footer.
		
		Returns:
			str: Text with qualifying repeated headers and footers removed.
		
		Raises:
			Error: Raised when validation, file access, or header and footer removal fails.
		"""
		try:
			throw_if( 'filepath', filepath )
			if not os.path.exists( filepath ):
				raise FileNotFoundError( f'File not found: {filepath}' )
			
			if lines < 6:
				raise ValueError( 'Argument "lines" must be at least 6.' )
			
			if headers < 0 or footers < 0:
				raise ValueError( 'Arguments "headers" and "footers" must be non-negative.' )
			
			if headers + footers >= lines:
				raise ValueError( 'The combined header and footer line counts must be less than '
				                  'the page line count.' )
			
			self.file_path = filepath
			with open( self.file_path, 'r', encoding='utf-8', errors='ignore' ) as file:
				self.lines = file.readlines( )
			
			self.pages = [ self.lines[ index:index + lines ] for index in
				range( 0, len( self.lines ), lines ) ]
			
			if len( self.pages ) < 2:
				return ''.join( self.lines )
			
			header_counts: Dict[ Tuple[ str, ... ], int ] = { }
			footer_counts: Dict[ Tuple[ str, ... ], int ] = { }
			header_page_count = 0
			footer_page_count = 0
			for page in self.pages:
				page_length = len( page )
				if headers > 0 and page_length >= headers:
					header_value = tuple( page[ :headers ] )
					header_counts[ header_value ] = (header_counts.get( header_value, 0 ) + 1)
					header_page_count += 1
				
				if footers > 0 and page_length >= footers:
					footer_value = tuple( page[ -footers: ] )
					footer_counts[ footer_value ] = (footer_counts.get( footer_value, 0 ) + 1)
					footer_page_count += 1
			
			common_header: Tuple[ str, ... ] = ( )
			common_footer: Tuple[ str, ... ] = ( )
			if header_counts:
				header_value, header_count = max( header_counts.items( ),
					key=lambda item: item[ 1 ] )
				
				if (header_count >= 2 and header_count / header_page_count >= 0.5):
					common_header = header_value
			
			if footer_counts:
				footer_value, footer_count = max( footer_counts.items( ),
					key=lambda item: item[ 1 ] )
				
				if (footer_count >= 2 and footer_count / footer_page_count >= 0.5):
					common_footer = footer_value
			
			cleaned_pages: List[ str ] = [ ]
			for page in self.pages:
				page_lines = list( page )
				if (common_header and len( page_lines ) >= len( common_header ) and tuple(
						page_lines[ :len( common_header ) ] ) == common_header):
					page_lines = page_lines[ len( common_header ): ]
				
				if (common_footer and len( page_lines ) >= len( common_footer ) and tuple(
						page_lines[ -len( common_footer ): ] ) == common_footer):
					page_lines = page_lines[ :-len( common_footer ) ]
				
				cleaned_pages.append( ''.join( page_lines ) )
			
			return ''.join( cleaned_pages )
		except Exception as e:
			exception = Error( e )
			exception.module = 'processors'
			exception.cause = 'TextParser'
			exception.method = ('remove_headers( self, filepath: str, lines: int=50, '
			                    'headers: int=3, footers: int=3 ) -> str')
			Logger( ).write( exception )
			raise exception
	
	def remove_numbers( self, text: str ) -> str | None:
		"""Remove numeric digits.
		
		Purpose:
			Removes contiguous decimal digit sequences while preserving capitalization,
			punctuation, surrounding text, and readable token boundaries.
		
		Args:
			text (str): Source text containing numeric digits to remove.
		
		Returns:
			str | None: Text with decimal digit sequences removed.
		
		Raises:
			Error: Raised when validation or numeric-digit removal fails.
		"""
		try:
			throw_if( 'text', text )
			_cleaned = re.sub( r'\d+', ' ', text )
			_cleaned = re.sub( r'[ \t]{2,}', ' ', _cleaned )
			return _cleaned.strip( )
		except Exception as e:
			exception = Error( e )
			exception.module = 'processors'
			exception.cause = 'TextParser'
			exception.method = 'remove_numbers( self, text: str ) -> str'
			Logger( ).write( exception )
			raise exception
	
	def remove_numerals( self, text: str ) -> str:
		"""Remove Roman numerals.
		
		Purpose:
			Removes standalone uppercase Roman numerals containing at least two characters while
			preserving single-letter words, lowercase text, ordinary numbers, punctuation,
			whitespace, citations, and document order.
		
		Args:
			text: Source text containing Roman numerals to remove.
		
		Returns:
			str: Text with qualifying standalone Roman numerals removed.
		
		Raises:
			Error: Raised when validation or Roman-numeral removal fails.
		"""
		try:
			throw_if( 'text', text )
			self.raw_input = text
			
			pattern = (r'\b(?=[MDCLXVI]{2,}\b)'
			           r'M{0,4}(?:CM|CD|D?C{0,3})'
			           r'(?:XC|XL|L?X{0,3})'
			           r'(?:IX|IV|V?I{0,3})\b')
			
			self.cleaned_text = re.sub( pattern, ' ', self.raw_input )
			self.cleaned_text = re.sub( r'[ \t]{2,}', ' ', self.cleaned_text )
			self.cleaned_text = re.sub( r'[ \t]+\n', '\n', self.cleaned_text )
			
			return self.cleaned_text.strip( )
		except Exception as e:
			exception = Error( e )
			exception.module = 'processors'
			exception.cause = 'TextParser'
			exception.method = 'remove_numerals( self, text: str ) -> str'
			Logger( ).write( exception )
			raise exception
	
	def remove_images( self, text: str ) -> str:
		"""Remove textual image references.
		
		Purpose:
			Removes Markdown image declarations, HTML image elements, data-image references, and
			standalone image URLs while preserving surrounding text, capitalization, document
			order, and non-image links.
		
		Args:
			text (str): Source text containing textual image references.
		
		Returns:
			str: Text with supported image references removed.
		
		Raises:
			Error: Raised when validation or image-reference removal fails.
		"""
		try:
			throw_if( 'text', text )
			self.raw_input = text
			_text = re.sub( r'!\[[^\]]*]\([^)]+\)', ' ', self.raw_input )
			_text = re.sub( r'!\[[^\]]*]\[[^\]]*]', ' ', _text )
			_text = re.sub( r'<img\b[^>]*>', ' ', _text, flags=re.IGNORECASE )
			_text = re.sub( r'data:image/[a-zA-Z0-9.+-]+;base64,[a-zA-Z0-9+/=\s]+', ' ', _text,
				flags=re.IGNORECASE )
			_text = re.sub( r'https?://[^\s<>"\']+\.(?:png|jpe?g|gif|bmp|svg|webp|tiff?|ico)'
			                r'(?:\?[^\s<>"\']*)?', ' ', _text, flags=re.IGNORECASE )
			_text = re.sub( r'[ \t]{2,}', ' ', _text )
			_text = re.sub( r' *\n *', '\n', _text )
			return _text.strip( )
		except Exception as e:
			exception = Error( e )
			exception.module = 'processors'
			exception.cause = 'TextParser'
			exception.method = 'remove_images( self, text: str ) -> str'
			Logger( ).write( exception )
			raise exception
	
	def tiktokenize( self, text: str, encoding: str='cl100k_base' ) -> DataFrame:
		"""Tokenize text with a TikToken encoding.
		
		Purpose:
			Encodes text into model-compatible token identifiers without altering capitalization,
			punctuation, numbers, whitespace, symbols, or content order.
		
		Args:
			text: Source text to tokenize.
			encoding: TikToken encoding name used to generate token identifiers.
		
		Returns:
			DataFrame: Token identifiers in source order.
		
		Raises:
			Error: Raised when validation, encoding resolution, or tokenization fails.
		"""
		try:
			throw_if( 'text', text )
			throw_if( 'encoding', encoding )			
			self.encoding = tiktoken.get_encoding( encoding )
			token_ids = self.encoding.encode( text )
			return pd.DataFrame( token_ids )
		except Exception as e:
			exception = Error( e )
			exception.module = 'processors'
			exception.cause = 'TextParser'
			exception.method = ('tiktokenize( self, text: str, '
			                    "encoding: str='cl100k_base' ) -> DataFrame")
			Logger( ).write( exception )
			raise exception
	
	def split_sentences( self, text: str ) -> List[ str ]:
		"""Split text into sentences.
		
		Purpose:
			Segments text into sentences without altering capitalization, punctuation, numbers,
			symbols, whitespace within sentences, legal citations, or content order.
		
		Args:
			text: Source text to segment into sentences.
		
		Returns:
			List[ str ]: Sentences in source order.
		
		Raises:
			Error: Raised when validation or sentence tokenization fails.
		"""
		try:
			throw_if( 'text', text )
			
			sentences = sent_tokenize( text )
			
			return [ sentence.strip( ) for sentence in sentences if sentence.strip( ) ]
		except Exception as e:
			exception = Error( e )
			exception.module = 'processors'
			exception.cause = 'TextParser'
			exception.method = ('split_sentences( self, text: str ) -> List[ str ]')
			Logger( ).write( exception )
			raise exception
	
	def split_pages( self, filepath: str, num: int=50 ) -> List[ str ]:
		"""Split text into logical pages.
		
		Purpose:
			Splits UTF-8 text into form-feed-delimited pages when page delimiters are available,
			or fixed-length line groups otherwise, while resetting parser state for each call.
		
		Args:
			filepath: Path to the UTF-8 text file to split.
			num: Maximum number of lines assigned to each fallback logical page.
		
		Returns:
			List[ str ]: Nonempty logical pages in source order.
		
		Raises:
			Error: Raised when validation, file access, or page splitting fails.
		"""
		try:
			throw_if( 'filepath', filepath )
			if not os.path.exists( filepath ):
				raise FileNotFoundError( f'File not found: {filepath}' )
			
			if num <= 0:
				raise ValueError( 'Argument "num" must be greater than zero.' )
			
			self.file_path = filepath
			self.pages = [ ]
			with open( self.file_path, 'r', encoding='utf-8', errors='ignore' ) as file:
				content = file.read( )
			
			if '\f' in content:
				self.pages = [ page.strip( ) for page in content.split( '\f' ) if page.strip( ) ]
				return self.pages
			
			self.lines = content.splitlines( )
			for index in range( 0, len( self.lines ), num ):
				page_lines = self.lines[ index:index + num ]
				page_text = '\n'.join( page_lines ).strip( )
				if page_text:
					self.pages.append( page_text )
			
			return self.pages
		except Exception as e:
			exception = Error( e )
			exception.module = 'processors'
			exception.cause = 'TextParser'
			exception.method = ('split_pages( self, filepath: str, num: int=50 ) '
			                    '-> List[ str ]')
			Logger( ).write( exception )
			raise exception
	
	def split_paragraphs( self, filepath: str ) -> DataFrame:
		"""Split text into paragraphs.
		
		Purpose:
			Splits a text file at blank-line paragraph boundaries while preserving internal
			sentence spacing, punctuation, capitalization, numbers, symbols, and paragraph order.
		
		Args:
			filepath: Path to the text file to split into paragraphs.
		
		Returns:
			DataFrame: Nonempty paragraphs in source order.
		
		Raises:
			Error: Raised when validation, file access, decoding, or paragraph splitting fails.
		"""
		try:
			throw_if( 'filepath', filepath )
			if not os.path.exists( filepath ):
				raise FileNotFoundError( f'File not found: {filepath}' )
			
			self.file_path = filepath
			self.paragraphs = [ ]
			
			try:
				with open( self.file_path, 'r', encoding='utf-8' ) as file:
					source_text = file.read( )
			except UnicodeDecodeError:
				with open( self.file_path, 'r', encoding='latin1' ) as file:
					source_text = file.read( )
			
			source_text = source_text.replace( '\r\n', '\n' )
			source_text = source_text.replace( '\r', '\n' )
			self.paragraphs = [ re.sub( r'[ \t]*\n[ \t]*', ' ', paragraph ).strip( ) for paragraph
				in re.split( r'\n[ \t]*\n+', source_text ) if paragraph.strip( ) ]
			
			return pd.DataFrame( { 'Paragraph': self.paragraphs } )
		except Exception as e:
			exception = Error( e )
			exception.module = 'processors'
			exception.cause = 'TextParser'
			exception.method = ('split_paragraphs( self, filepath: str ) -> DataFrame')
			Logger( ).write( exception )
			raise exception
	
	def create_frequency_distribution( self, tokens: List[ str ] ) -> DataFrame | None:
		"""Create frequency distribution.
		
		Purpose:
			Executes the ``create frequency distribution`` operation for the ``TextParser``
			workflow, updating instance state where required and returning the processed result
			used by downstream Chonky loading, processing, analysis, tokenization, or embedding
			steps.
		
		Args:
			tokens: Tokens value used by the processing operation. Expected type: ``List[ str ]``.
		
		Returns:
			DataFrame | None: Result produced by the processing operation.
		
		Raises:
			Error: Raised when validation or processing fails.
		"""
		try:
			throw_if( 'tokens', tokens )
			self.tokens = tokens
			_freqdist = FreqDist( dict( Counter( self.tokens ) ) )
			_words = _freqdist.items( )
			_data = pd.DataFrame( _words, columns=[ 'Word', 'Frequency' ] )
			_data.index.name = 'ID'
			return _data
		except Exception as e:
			exception = Error( e )
			exception.module = 'processors'
			exception.cause = 'TextParser'
			exception.method = ('create_frequency_distribution(self, tokens: List[ str '
			                    '])->DataFrame')
			Logger( ).write( exception )
			raise exception
	
	def create_vocabulary( self, tokens: List[ str ] ) -> Series | None:
		"""Create vocabulary.
		
		Purpose:
			Executes the ``create vocabulary`` operation for the ``TextParser`` workflow,
			updating instance state where required and returning the processed result used by
			downstream Chonky loading, processing, analysis, tokenization, or embedding steps.
		
		Args:
			tokens: Tokens value used by the processing operation. Expected type: ``List[ str ]``.
		
		Returns:
			Series | None: Result produced by the processing operation.
		
		Raises:
			Error: Raised when validation or processing fails.
		"""
		try:
			throw_if( 'tokens', tokens )
			self.tokens = tokens
			_freqdist = FreqDist( dict( Counter( self.tokens ) ) )
			_vocab = _freqdist.items( )
			_vocabulary = pd.DataFrame( _vocab, columns=[ 'Word', 'Frequency' ] )
			_words = _vocabulary.iloc[ :, 0 ]
			return _words
		except Exception as e:
			exception = Error( e )
			exception.module = 'processors'
			exception.cause = 'TextParser'
			exception.method = ('create_vocabulary(self, freq_dist: dict, min: int=1)->List[str]')
			Logger( ).write( exception )
			raise exception
	
	def create_wordbag( self, tokens: List[ str ] ) -> DataFrame | None:
		"""Create wordbag.
		
		Purpose:
			Executes the ``create wordbag`` operation for the ``TextParser`` workflow, updating
			instance state where required and returning the processed result used by downstream
			Chonky loading, processing, analysis, tokenization, or embedding steps.
		
		Args:
			tokens: Tokens value used by the processing operation. Expected type: ``List[ str ]``.
		
		Returns:
			DataFrame | None: Result produced by the processing operation.
		
		Raises:
			Error: Raised when validation or processing fails.
		"""
		try:
			throw_if( 'tokens', tokens )
			self.tokens = tokens
			_freqdist = FreqDist( dict( Counter( self.tokens ) ) )
			_words = _freqdist.keys( )
			_data = pd.DataFrame( _words )
			return _data
		except Exception as e:
			exception = Error( e )
			exception.module = 'processors'
			exception.cause = 'TextParser'
			exception.method = 'create_wordbag( self, words: List[ str ] ) -> dict'
			Logger( ).write( exception )
			raise exception
	
	def create_vectors( self, tokens: List[ str ] ) -> DataFrame:
		"""Create TF-IDF vectors.
		
		Purpose:
			Creates one TF-IDF feature vector for each supplied token while preserving token
			capitalization, punctuation, symbols, duplicates, and source order.
		
		Args:
			tokens: Tokens used to construct the TF-IDF feature matrix.
		
		Returns:
			DataFrame: TF-IDF vectors indexed by the source tokens.
		
		Raises:
			Error: Raised when validation or TF-IDF vector creation fails.
		"""
		try:
			throw_if( 'tokens', tokens )
			
			if not tokens:
				raise ValueError( 'Argument "tokens" must contain at least one token.' )
			
			if any( not isinstance( token, str ) or not token for token in tokens ):
				raise ValueError( 'Argument "tokens" must contain only nonempty strings.' )
			
			self.tokens = tokens
			self.vectorizer = TfidfVectorizer( tokenizer=str.split, preprocessor=None,
				token_pattern=None, lowercase=False )
			
			vector_matrix = self.vectorizer.fit_transform( self.tokens )
			feature_names = self.vectorizer.get_feature_names_out( )
			
			return pd.DataFrame( vector_matrix.toarray( ), columns=feature_names,
				index=pd.Index( self.tokens, name='Token' ) )
		except Exception as e:
			exception = Error( e )
			exception.module = 'processors'
			exception.cause = 'TextParser'
			exception.method = ('create_vectors( self, tokens: List[ str ] ) '
			                    '-> DataFrame')
			Logger( ).write( exception )
			raise exception
	
	def clean_file( self, filepath: str ) -> str | None:
		"""Clean file.
		
		Purpose:
			Executes the ``clean file`` operation for the ``TextParser`` workflow, updating
			instance state where required and returning the processed result used by downstream
			Chonky loading, processing, analysis, tokenization, or embedding steps.
		
		Args:
			filepath: Filepath value used by the processing operation. Expected type: ``str``.
		
		Returns:
			str | None: Result produced by the processing operation.
		
		Raises:
			Error: Raised when validation or processing fails.
		"""
		try:
			throw_if( 'filepath', filepath )
			if not os.path.exists( filepath ):
				raise FileNotFoundError( f'File not found: {filepath}' )
			else:
				_sourcepath = filepath
				_text = open( _sourcepath, 'r', encoding='utf-8', errors='ignore' ).read( )
				_collapsed = self.collapse_whitespace( _text )
				_compressed = self.compress_whitespace( _collapsed )
				_normalized = self.normalize_text( _compressed )
				_encoded = self.remove_encodings( _normalized )
				_special = self.remove_symbols( _encoded )
				_cleaned = self.remove_fragments( _special )
				_recompress = self.compress_whitespace( _cleaned )
				_lemmatized = self.lemmatize_text( _recompress )
				_stops = self.remove_stopwords( _lemmatized )
				return _stops
		except Exception as e:
			exception = Error( e )
			exception.module = 'processors'
			exception.cause = 'TextParser'
			exception.method = 'clean_file( self, src: str ) -> str'
			Logger( ).write( exception )
			raise exception
	
	def clean_files( self, source: str, destination: str ) -> None:
		"""Clean files.
		
		Purpose:
			Executes the ``clean files`` operation for the ``TextParser`` workflow, updating
			instance state where required and returning the processed result used by downstream
			Chonky loading, processing, analysis, tokenization, or embedding steps.
		
		Args:
			source: Source value used by the processing operation. Expected type: ``str``.
			destination: Destination value used by the processing operation. Expected type:
			``str``.
		
		Returns:
			None: Result produced by the processing operation.
		
		Raises:
			Error: Raised when validation or processing fails.
		"""
		try:
			throw_if( 'src', source )
			throw_if( 'dest', destination )
			if not os.path.exists( source ):
				raise FileNotFoundError( f'File not found: {source}' )
			elif not os.path.exists( destination ):
				raise FileNotFoundError( f'File not found: {destination}' )
			else:
				_source = source
				_destpath = destination
				_files = os.listdir( _source )
				for f in _files:
					_processed = [ ]
					_filename = os.path.basename( f )
					_sourcepath = _source + '\\' + _filename
					_text = open( _sourcepath, 'r', encoding='utf-8', errors='ignore' ).read( )
					_collapsed = self.collapse_whitespace( _text )
					_compressed = self.compress_whitespace( _collapsed )
					_normalized = self.normalize_text( _compressed )
					_encoded = self.remove_encodings( _normalized )
					_special = self.remove_symbols( _encoded )
					_cleaned = self.remove_fragments( _special )
					_recompress = self.compress_whitespace( _cleaned )
					_lemmatized = self.lemmatize_text( _recompress )
					_stops = self.remove_stopwords( _lemmatized )
					_sentences = self.split_sentences( _stops )
					_destination = _destpath + '\\' + _filename
					_clean = open( _destination, 'wt', encoding='utf-8', errors='ignore' )
					_lines = ' '.join( _sentences )
					_clean.write( _lines )
		except Exception as e:
			exception = Error( e )
			exception.module = 'processors'
			exception.cause = 'TextParser'
			exception.method = 'clean_files( self, src: str, dest: str )'
			Logger( ).write( exception )
			raise exception
	
	def chunk_files( self, source: str, destination: str ) -> None:
		"""Chunk files.
		
		Purpose:
			Executes the ``chunk files`` operation for the ``TextParser`` workflow, updating
			instance state where required and returning the processed result used by downstream
			Chonky loading, processing, analysis, tokenization, or embedding steps.
		
		Args:
			source: Source value used by the processing operation. Expected type: ``str``.
			destination: Destination value used by the processing operation. Expected type:
			``str``.
		
		Returns:
			None: Result produced by the processing operation.
		
		Raises:
			Error: Raised when validation or processing fails.
		"""
		try:
			throw_if( 'src', source )
			throw_if( 'dest', destination )
			if not os.path.exists( source ):
				raise FileNotFoundError( f'File not found: {source}' )
			elif not os.path.exists( destination ):
				raise FileNotFoundError( f'File not found: {destination}' )
			else:
				_source = source
				_destination = destination
				_files = os.listdir( _source )
				_words = [ ]
				for f in _files:
					_processed = [ ]
					_filename = os.path.basename( f )
					_sourcepath = _source + '\\' + _filename
					_text = open( _sourcepath, 'r', encoding='utf-8', errors='ignore' ).read( )
					_sentences = self.split_sentences( _text )
					_datamap = [ ]
					for v in _sentences:
						_datamap.append( v )
					
					for s in _datamap:
						_processed.append( s )
					
					_final = _destination + '\\' + _filename
					_clean = open( _final, 'wt', encoding='utf-8', errors='ignore' )
					for p in _processed:
						_clean.write( p )
		except Exception as e:
			exception = Error( e )
			exception.module = 'processors'
			exception.cause = 'TextParser'
			exception.method = 'chunk_files( self, src: str, dest: str )'
			Logger( ).write( exception )
			raise exception
	
	def chunk_data( self, filepath: str, size: int=10 ) -> DataFrame | None:
		"""Chunk data.
		
		Purpose:
			Executes the ``chunk data`` operation for the ``TextParser`` workflow, updating
			instance state where required and returning the processed result used by downstream
			Chonky loading, processing, analysis, tokenization, or embedding steps.
		
		Args:
			filepath: Filepath value used by the processing operation. Expected type: ``str``.
			size: Size value used by the processing operation. Expected type: ``int``. Defaults to
			``10``.
		
		Returns:
			DataFrame | None: Result produced by the processing operation.
		
		Raises:
			Error: Raised when validation or processing fails.
		"""
		try:
			throw_if( 'filepath', filepath )
			if not os.path.exists( filepath ):
				raise FileNotFoundError( f'File not found: {filepath}' )
			else:
				_source = filepath
				_processed = [ ]
				_wordlist = [ ]
				_vocab = words.words( 'en' )
				_text = open( _source, 'r', encoding='utf-8', errors='ignore' ).read( )
				_lower = _text.lower( )
				_tokens = _lower.split( )
				for s in _tokens:
					if s.isalpha( ) and s in _vocab:
						_wordlist.append( s )
				self.chunks = [ _wordlist[ i: i + size ] for i in
					range( 0, len( _wordlist ), size ) ]
				for i, c in enumerate( self.chunks ):
					_item = '[' + ' '.join( c ) + '],'
					_processed.append( _item )
				_data = pd.DataFrame( _processed )
				return _data
		except Exception as e:
			exception = Error( e )
			exception.module = 'processors'
			exception.cause = 'TextParser'
			exception.method = 'chunk_data( self, filepath: str, size: int=512  ) -> DataFrame'
			Logger( ).write( exception )
			raise exception
	
	def chunk_datasets( self, source: str, destination: str, size: int=10 ) -> DataFrame:
		"""Chunk datasets.
		
		Purpose:
			Executes the ``chunk datasets`` operation for the ``TextParser`` workflow, updating
			instance state where required and returning the processed result used by downstream
			Chonky loading, processing, analysis, tokenization, or embedding steps.
		
		Args:
			source: Source value used by the processing operation. Expected type: ``str``.
			destination: Destination value used by the processing operation. Expected type:
			``str``.
			size: Size value used by the processing operation. Expected type: ``int``. Defaults to
			``10``.
		
		Returns:
			DataFrame: Result produced by the processing operation.
		
		Raises:
			Error: Raised when validation or processing fails.
		"""
		try:
			throw_if( 'filepath', source )
			throw_if( 'destination', destination )
			if not os.path.exists( source ):
				raise FileNotFoundError( f'File not found: {source}' )
			elif not os.path.exists( destination ):
				raise FileNotFoundError( f'File not found: {destination}' )
			else:
				_src = source
				_destination = destination
				_files = os.listdir( _src )
				_words = [ ]
				for f in _files:
					_processed = [ ]
					_filename = os.path.basename( f )
					_sourcepath = _src + '\\' + _filename
					_text = open( _sourcepath, 'r', encoding='utf-8', errors='ignore' ).read( )
					_collapsed = self.collapse_whitespace( _text )
					_compressed = self.compress_whitespace( _collapsed )
					_normalized = self.normalize_text( _compressed )
					_encoded = self.remove_encodings( _normalized )
					_special = self.remove_symbols( _encoded )
					_cleaned = self.remove_fragments( _special )
					_recompress = self.compress_whitespace( _cleaned )
					_lemmatized = self.lemmatize_text( _recompress )
					_stops = self.remove_stopwords( _lemmatized )
					_tokens = _stops.split( None )
					_chunks = [ _tokens[ i: i + size ] for i in range( 0, len( _tokens ), size ) ]
					_datamap = [ ]
					for i, c in enumerate( _chunks ):
						_row = ' '.join( c )
						_datamap.append( _row )
					
					for s in _datamap:
						_processed.append( s )
					
					_name = _filename.replace( '.txt', '.xlsx' )
					_savepath = (_destination + f'\\' + _name)
					_data = pd.DataFrame( _processed, columns=[ 'Data', ] )
					_data.to_excel( _savepath, sheet_name='Dataset', index=False,
						columns=[ 'Data', ] )
		except Exception as e:
			exception = Error( e )
			exception.module = 'processors'
			exception.cause = 'TextParser'
			exception.method = 'chunk_data( self, filepath: str, size: int=15  ) -> DataFrame'
			Logger( ).write( exception )
			raise exception
	
	def convert_jsonl( self, source: str, destination: str, size: int=10 ) -> None:
		"""Convert jsonl.
		
		Purpose:
			Executes the ``convert jsonl`` operation for the ``TextParser`` workflow, updating
			instance state where required and returning the processed result used by downstream
			Chonky loading, processing, analysis, tokenization, or embedding steps.
		
		Args:
			source: Source value used by the processing operation. Expected type: ``str``.
			destination: Destination value used by the processing operation. Expected type:
			``str``.
			size: Size value used by the processing operation. Expected type: ``int``. Defaults to
			``10``.
		
		Returns:
			None: Result produced by the processing operation.
		
		Raises:
			Error: Raised when validation or processing fails.
		"""
		try:
			throw_if( 'source', source )
			throw_if( 'destination', destination )
			if not os.path.exists( source ):
				raise FileNotFoundError( f'File not found: {source}' )
			elif not os.path.exists( destination ):
				raise FileNotFoundError( f'File not found: {destination}' )
			else:
				_source = source
				_destpath = destination
				_files = os.listdir( _source )
				_wordlist = [ ]
				for f in _files:
					_processed = [ ]
					_filename = os.path.basename( f )
					_sourcepath = _source + '\\' + _filename
					_text = open( _sourcepath, 'r', encoding='utf-8', errors='ignore' ).read( )
					_tokens = _text.split( ' ' )
					_chunks = [ _tokens[ i: i + size ] for i in range( 0, len( _tokens ), size ) ]
					_datamap = [ ]
					for i, c in enumerate( _chunks ):
						_value = '{ ' + f' {i} : [ ' + ' '.join( c ) + ' ] }, ' + "\n"
						_datamap.append( _value )
					
					for s in _datamap:
						_processed.append( s )
					
					_destination = _destpath + '\\' + _filename.replace( '.txt', '.jsonl' )
					_clean = open( _destination, 'wt', encoding='utf-8', errors='ignore' )
					for p in _processed:
						_clean.write( p )
		except Exception as e:
			exception = Error( e )
			exception.module = 'processors'
			exception.cause = 'TextParser'
			exception.method = 'convert_jsonl( self, source: str, desination: str )'
			Logger( ).write( exception )
			raise exception
	
	def encode_sentences( self, tokens: List[ str ], model: str='all-MiniLM-L6-v2' ) -> Tuple[ List[ str ], np.ndarray ]:
		"""Encode sentences.
		
		Purpose:
			Executes the ``encode sentences`` operation for the ``TextParser`` workflow, updating
			instance state where required and returning the processed result used by downstream
			Chonky loading, processing, analysis, tokenization, or embedding steps.
		
		Args:
			tokens: Tokens value used by the processing operation. Expected type: ``List[ str ]``.
			model: Model value used by the processing operation. Expected type: ``str``. Defaults
			to ``'all-MiniLM-L6-v2'``.
		
		Returns:
			Tuple[ List[ str ], np.ndarray ]: Result produced by the processing operation.
		
		Raises:
			Error: Raised when validation or processing fails.
		"""
		try:
			throw_if( 'tokens', tokens )
			throw_if( 'model', model )
			_transformer = SentenceTransformer( model )
			_tokens = [ self.lemmatizer.lemmatize( t ) for t in tokens ]
			_encoding = _transformer.encode( _tokens, show_progress_bar=True )
			return (self.cleaned_tokens, np.array( _encoding ))
		except Exception as e:
			exception = Error( e )
			exception.module = 'processors'
			exception.cause = 'TextParser'
			exception.method = ('encode_sentences( self, sentences: List[ str ], model_name ) -> '
			                    '( )')
			Logger( ).write( exception )
			raise exception
	
	def semantic_search( self, query: str, tokens: List[ str ], embeddings: np.ndarray,
		model: SentenceTransformer, top: int=5 ) -> List[ tuple[ str, float ] ]:
		"""Semantic search.
		
		Purpose:
			Executes the ``semantic search`` operation for the ``TextParser`` workflow, updating
			instance state where required and returning the processed result used by downstream
			Chonky loading, processing, analysis, tokenization, or embedding steps.
		
		Args:
			query: Query value used by the processing operation. Expected type: ``str``.
			tokens: Tokens value used by the processing operation. Expected type: ``List[ str ]``.
			embeddings: Embeddings value used by the processing operation. Expected type:
			``np.ndarray``.
			model: Model value used by the processing operation. Expected type:
			``SentenceTransformer``.
			top: Top value used by the processing operation. Expected type: ``int``. Defaults to
			``5``.
		
		Returns:
			List[ tuple[ str, float ] ]: Result produced by the processing operation.
		
		Raises:
			Error: Raised when validation or processing fails.
		"""
		try:
			throw_if( 'query', query )
			throw_if( 'tokens', tokens )
			throw_if( 'embedding', embeddings )
			throw_if( 'model', model )
			query_vec = model.encode( [ query ] )
			sims = cosine_similarity( query_vec, embeddings )[ 0 ]
			top_indices = sims.argsort( )[ ::-1 ][ : top ]
			return [ (tokens[ i ], sims[ i ]) for i in top_indices ]
		except Exception as e:
			exception = Error( e )
			exception.module = 'processors'
			exception.cause = 'TextParser'
			exception.method = ('semantic_search( self, query: str, tokens: List[ str ], '
			                    'embeddings: np.ndarray, model: SentenceTransformer,  '
			                    'top_k: int=5 ) -> List[ tuple[ str, float ] ]')
			Logger( ).write( exception )
			raise exception

class NltkParser( Processor ):
	"""NltkParser processing component.
	
	Purpose:
		Provides NLTK-backed parsing utilities for corpus tokenization, stemming, lemmatization,
		part-of-speech tagging, named-entity handling, and lexical diagnostics used by Chonky
		analysis workflows.
	
	Attributes:
		word_tokens: Runtime state used by ``NltkParser`` during Chonky processing operations.
		sentence_tokens: Runtime state used by ``NltkParser`` during Chonky processing operations.
		stemmed_tokens: Runtime state used by ``NltkParser`` during Chonky processing operations.
		lemmatized_tokens: Runtime state used by ``NltkParser`` during Chonky processing
		operations.
		tagged_tokens: Runtime state used by ``NltkParser`` during Chonky processing operations.
		named_entities: Runtime state used by ``NltkParser`` during Chonky processing operations.
	"""
	word_tokens: Optional[ List[ str ] ]
	sentence_tokens: Optional[ List[ str ] ]
	stemmed_tokens: Optional[ List[ str ] ]
	lemmatized_tokens: Optional[ List[ str ] ]
	tagged_tokens: Optional[ List[ Tuple[ str, str ] ] ]
	named_entities: Optional[ List[ Tuple[ str, str ] ] ]
	
	def __init__( self ) -> None:
		"""Initialize the NltkParser instance.
		
		Purpose:
			Initializes parser state, reusable helper objects, and runtime caches used by later
			processing methods.
		"""
		super( ).__init__( )
		self.initialize_resources( )
		self.word_tokens = [ ]
		self.sentence_tokens = [ ]
		self.stemmed_tokens = [ ]
		self.lemmatized_tokens = [ ]
		self.tagged_tokens = [ ]
		self.named_entities = [ ]
	
	def __dir__( self ) -> List[ str ] | None:
		"""Dir.
		
		Purpose:
			Returns the public member names exposed by the parser for introspection, diagnostics,
			and MkDocs API documentation.
		
		Returns:
			List[ str ] | None: Result produced by the processing operation.
		"""
		return [ 'initialize_resources', 'word_tokenizer', 'sentence_tokenizer', 'word_stemmer',
			'word_lemmatizer', 'pos_tagger', 'named_entity_recognition', 'word_tokens',
			'sentence_tokens', 'stemmed_tokens', 'lemmatized_tokens', 'tagged_tokens',
			'named_entities' ]
	
	def initialize_resources( self ) -> None:
		"""Initialize resources.
		
		Purpose:
			Executes the ``initialize resources`` operation for the ``NltkParser`` workflow,
			updating instance state where required and returning the processed result used by
			downstream Chonky loading, processing, analysis, tokenization, or embedding steps.
		
		Returns:
			None: Result produced by the processing operation.
		
		Raises:
			Error: Raised when validation or processing fails.
		"""
		try:
			required_resources: List[ Tuple[ str, str ] ] = [ ('tokenizers/punkt', 'punkt'),
				('tokenizers/punkt_tab', 'punkt_tab'), ('corpora/wordnet', 'wordnet'),
				('corpora/omw-1.4', 'omw-1.4'),
				('taggers/averaged_perceptron_tagger', 'averaged_perceptron_tagger'),
				('taggers/averaged_perceptron_tagger_eng', 'averaged_perceptron_tagger_eng'),
				('chunkers/maxent_ne_chunker', 'maxent_ne_chunker'),
				('chunkers/maxent_ne_chunker_tab', 'maxent_ne_chunker_tab'),
				('corpora/words', 'words'), ]
			
			for resource_path, resource_name in required_resources:
				try:
					nltk.data.find( resource_path )
				except LookupError:
					nltk.download( resource_name )
		except Exception as e:
			exception = Error( e )
			exception.module = 'processing'
			exception.cause = 'NltkParser'
			exception.method = 'NltkParser._ensure_nltk_resources( self ) -> None'
			Logger( ).write( exception )
			raise exception
	
	def word_tokenizer( self, text: str ) -> List[ str ] | None:
		"""Word tokenizer.
		
		Purpose:
			Executes the ``word tokenizer`` operation for the ``NltkParser`` workflow, updating
			instance state where required and returning the processed result used by downstream
			Chonky loading, processing, analysis, tokenization, or embedding steps.
		
		Args:
			text: Text value used by the processing operation. Expected type: ``str``.
		
		Returns:
			List[ str ] | None: Result produced by the processing operation.
		
		Raises:
			Error: Raised when validation or processing fails.
		"""
		try:
			throw_if( 'text', text )
			_text = text.lower( )
			self.word_tokens = word_tokenize( _text )
			words = [ token for token in self.word_tokens ]
			return words
		except Exception as e:
			exception = Error( e )
			exception.module = 'processing'
			exception.cause = 'NltkParser'
			exception.method = 'word_tokenizer( self, text: str ) -> List[ str ]'
			Logger( ).write( exception )
			raise exception
	
	def sentence_tokenizer( self, text: str ) -> List[ str ] | None:
		"""Sentence tokenizer.
		
		Purpose:
			Executes the ``sentence tokenizer`` operation for the ``NltkParser`` workflow,
			updating instance state where required and returning the processed result used by
			downstream Chonky loading, processing, analysis, tokenization, or embedding steps.
		
		Args:
			text: Text value used by the processing operation. Expected type: ``str``.
		
		Returns:
			List[ str ] | None: Result produced by the processing operation.
		
		Raises:
			Error: Raised when validation or processing fails.
		"""
		try:
			throw_if( 'text', text )
			_text = text.lower( )
			self.sentence_tokens = sent_tokenize( _text )
			return self.sentence_tokens
		except Exception as e:
			exception = Error( e )
			exception.module = 'processing'
			exception.cause = 'NltkParser'
			exception.method = 'tokenize_sentences( self, text: str ) -> str'
			Logger( ).write( exception )
			raise exception
	
	def word_stemmer( self, text: str ) -> List[ str ] | None:
		"""Word stemmer.
		
		Purpose:
			Executes the ``word stemmer`` operation for the ``NltkParser`` workflow, updating
			instance state where required and returning the processed result used by downstream
			Chonky loading, processing, analysis, tokenization, or embedding steps.
		
		Args:
			text: Text value used by the processing operation. Expected type: ``str``.
		
		Returns:
			List[ str ] | None: Result produced by the processing operation.
		
		Raises:
			Error: Raised when validation or processing fails.
		"""
		try:
			throw_if( 'text', text )
			_text = text.lower( )
			self.word_tokens = word_tokenize( _text )
			self.stemmed_tokens = [ self.stemmer.stem( t ) for t in self.word_tokens if
				isinstance( t, str ) and t.strip( ) ]
			
			return self.stemmed_tokens
		except Exception as e:
			exception = Error( e )
			exception.module = 'processing'
			exception.cause = 'NltkParser'
			exception.method = 'stemmer( self, text: str ) -> str'
			Logger( ).write( exception )
			raise exception
	
	def word_lemmatizer( self, text: str ) -> List[ str ] | None:
		"""Word lemmatizer.
		
		Purpose:
			Executes the ``word lemmatizer`` operation for the ``NltkParser`` workflow, updating
			instance state where required and returning the processed result used by downstream
			Chonky loading, processing, analysis, tokenization, or embedding steps.
		
		Args:
			text: Text value used by the processing operation. Expected type: ``str``.
		
		Returns:
			List[ str ] | None: Result produced by the processing operation.
		
		Raises:
			Error: Raised when validation or processing fails.
		"""
		try:
			throw_if( 'text', text )
			_text = text.lower( )
			self.word_tokens = word_tokenize( _text )
			self.lemmatized_tokens = [ self.lemmatizer.lemmatize( t ) for t in self.word_tokens if
				isinstance( t, str ) and t.strip( ) ]
			
			return self.lemmatized_tokens
		except Exception as e:
			exception = Error( e )
			exception.module = 'processing'
			exception.cause = 'NltkParser'
			exception.method = 'lemmatizer( self, text: str ) -> str'
			Logger( ).write( exception )
			raise exception
	
	def pos_tagger( self, text: str ) -> List[ Tuple[ str, str ] ] | None:
		"""Pos tagger.
		
		Purpose:
			Executes the ``pos tagger`` operation for the ``NltkParser`` workflow, updating
			instance state where required and returning the processed result used by downstream
			Chonky loading, processing, analysis, tokenization, or embedding steps.
		
		Args:
			text: Text value used by the processing operation. Expected type: ``str``.
		
		Returns:
			List[ Tuple[ str, str ] ] | None: Result produced by the processing operation.
		
		Raises:
			Error: Raised when validation or processing fails.
		"""
		try:
			throw_if( 'text', text )
			_text = text.lower( )
			self.word_tokens = word_tokenize( _text )
			self.tagged_tokens = nltk.pos_tag( self.word_tokens )
			return self.tagged_tokens
		except Exception as e:
			exception = Error( e )
			exception.module = 'processing'
			exception.cause = 'NltkParser'
			exception.method = 'pos_tagger( self, text: str ) -> str'
			Logger( ).write( exception )
			raise exception
	
	def named_entity_recognition( self, text: str ) -> List[ Tuple[ str, str ] ] | None:
		"""Named entity recognition.
		
		Purpose:
			Executes the ``named entity recognition`` operation for the ``NltkParser`` workflow,
			updating instance state where required and returning the processed result used by
			downstream Chonky loading, processing, analysis, tokenization, or embedding steps.
		
		Args:
			text: Text value used by the processing operation. Expected type: ``str``.
		
		Returns:
			List[ Tuple[ str, str ] ] | None: Result produced by the processing operation.
		
		Raises:
			Error: Raised when validation or processing fails.
		"""
		try:
			throw_if( 'text', text )
			_text = text.lower( )
			self.word_tokens = word_tokenize( _text )
			self.tagged_tokens = nltk.pos_tag( self.word_tokens )
			tree = nltk.ne_chunk( self.tagged_tokens )
			self.named_entities = [ ]
			for node in tree:
				if hasattr( node, 'label' ):
					label = node.label( )
					entity_text = ' '.join( token for token, _ in node.leaves( ) if
						isinstance( token, str ) and token.strip( ) )
					
					if entity_text:
						self.named_entities.append( (entity_text, label) )
			
			return self.named_entities
		except Exception as e:
			exception = Error( e )
			exception.module = 'processing'
			exception.cause = 'NltkParser'
			exception.method = 'named_entity_recogniztion( self, text: str ) -> str'
			Logger( ).write( exception )
			raise exception
	
	def chunk_words( self, text: str, size: int=5 ) -> DataFrame | None:
		"""Chunk words.
		
		Purpose:
			Executes the ``chunk words`` operation for the ``NltkParser`` workflow, updating
			instance state where required and returning the processed result used by downstream
			Chonky loading, processing, analysis, tokenization, or embedding steps.
		
		Args:
			text: Text value used by the processing operation. Expected type: ``str``.
			size: Size value used by the processing operation. Expected type: ``int``. Defaults to
			``5``.
		
		Returns:
			DataFrame | None: Result produced by the processing operation.
		
		Raises:
			Error: Raised when validation or processing fails.
		"""
		try:
			throw_if( 'text', text )
			_text = text.lower( )
			_tokens = nltk.word_tokenize( _text )
			_sentences = [ _tokens[ i: i + size ] for i in range( 0, len( _tokens ), size ) ]
			_datamap = [ ]
			for index, chunk in enumerate( _sentences ):
				_item = ' '.join( chunk )
				_datamap.append( _item )
			
			_data = pd.DataFrame( _datamap )
			return _data
		except Exception as e:
			exception = Error( e )
			exception.module = 'processors'
			exception.cause = 'TextParser'
			exception.method = 'chunk_sentences( self, text: str, max: int=10 ) -> DataFrame'
			Logger( ).write( exception )
			raise exception
	
	def chunk_sentences( self, text: str, size: int=15 ) -> DataFrame | None:
		"""Chunk sentences.
		
		Purpose:
			Executes the ``chunk sentences`` operation for the ``NltkParser`` workflow, updating
			instance state where required and returning the processed result used by downstream
			Chonky loading, processing, analysis, tokenization, or embedding steps.
		
		Args:
			text: Text value used by the processing operation. Expected type: ``str``.
			size: Size value used by the processing operation. Expected type: ``int``. Defaults to
			``15``.
		
		Returns:
			DataFrame | None: Result produced by the processing operation.
		
		Raises:
			Error: Raised when validation or processing fails.
		"""
		try:
			throw_if( 'text', text )
			_text = text.lower( )
			_tokens = sent_tokenize( _text )
			_sentences = [ _tokens[ i: i + size ] for i in range( 0, len( _tokens ), size ) ]
			_datamap = [ ]
			for i, c in enumerate( _sentences ):
				_item = ' '.join( c )
				_datamap.append( _item )
			
			_data = pd.DataFrame( _datamap )
			return _data
		except Exception as e:
			exception = Error( e )
			exception.module = 'processors'
			exception.cause = 'NltkParser'
			exception.method = 'chunk_sentences( self, text: str, max: int=512 ) -> DataFrame'
			Logger( ).write( exception )
			raise exception

class WordParser( Processor ):
	"""WordParser processing component.
	
	Purpose:
		Provides Microsoft Word document parsing utilities that extract paragraphs, tables,
		metadata, and cleaned text for downstream Chonky processing and analysis.
	
	Attributes:
		sentences: Runtime state used by ``WordParser`` during Chonky processing operations.
		cleaned_sentences: Runtime state used by ``WordParser`` during Chonky processing
		operations.
		document: Runtime state used by ``WordParser`` during Chonky processing operations.
		raw_text: Runtime state used by ``WordParser`` during Chonky processing operations.
		paragraphs: Runtime state used by ``WordParser`` during Chonky processing operations.
		file_path: Runtime state used by ``WordParser`` during Chonky processing operations.
		vocabulary: Runtime state used by ``WordParser`` during Chonky processing operations.
		document: Runtime state used by ``WordParser`` during Chonky processing operations.
	"""
	sentences: Optional[ List[ str ] ]
	cleaned_sentences: Optional[ List[ str ] ]
	document: Optional[ Document ]
	raw_text: Optional[ str ]
	paragraphs: Optional[ List[ str ] ]
	file_path: Optional[ str ]
	vocabulary: Optional[ set ]
	document: Optional[ Document ]
	
	def __init__( self, filepath: str ) -> None:
		"""Initialize the WordParser instance.
		
		Purpose:
			Initializes parser state, reusable helper objects, and runtime caches used by later
			processing methods.
		
		Args:
			filepath: Filepath value used by the processing operation. Expected type: ``str``.
		"""
		super( ).__init__( )
		self.file_path = filepath
		self.document = Document( open( self.file_path, 'rt+' ) )
		self.page_text = ''
		self.paragraphs = [ ]
		self.sentences = [ ]
		self.cleaned_sentences = [ ]
		self.vocabulary = set( )
	
	def __dir__( self ) -> List[ str ] | None:
		"""Dir.
		
		Purpose:
			Returns the public member names exposed by the parser for introspection, diagnostics,
			and MkDocs API documentation.
		
		Returns:
			List[ str ] | None: Result produced by the processing operation.
		"""
		return [ 'extract_text', 'split_sentences', 'clean_sentences', 'create_vocabulary',
			'compute_frequency_distribution', 'summarize', 'filepath', 'raw_text', 'paragraphs',
			'sentences', 'cleaned_sentences', 'vocabulary', 'freq_dist' ]
	
	def extract_text( self, num: int=1 ) -> str | None:
		"""Extract text.
		
		Purpose:
			Executes the ``extract text`` operation for the ``WordParser`` workflow, updating
			instance state where required and returning the processed result used by downstream
			Chonky loading, processing, analysis, tokenization, or embedding steps.
		
		Args:
			num: Num value used by the processing operation. Expected type: ``int``. Defaults to
			``1``.
		
		Returns:
			str | None: Result produced by the processing operation.
		
		Raises:
			Error: Raised when validation or processing fails.
		"""
		try:
			self.page_text = self.document.get_page_text( pno=num )
			return self.page_text
		except Exception as e:
			exception = Error( e )
			exception.module = 'processors'
			exception.cause = 'Word'
			exception.method = 'extract_text( self, num: int ) -> str'
			Logger( ).write( exception )
			raise exception
	
	def split_sentences( self ) -> List[ str ] | None:
		"""Split sentences.
		
		Purpose:
			Executes the ``split sentences`` operation for the ``WordParser`` workflow, updating
			instance state where required and returning the processed result used by downstream
			Chonky loading, processing, analysis, tokenization, or embedding steps.
		
		Returns:
			List[ str ] | None: Result produced by the processing operation.
		
		Raises:
			Error: Raised when validation or processing fails.
		"""
		try:
			_text = self.page_text.lower( )
			self.sentences = sent_tokenize( _text )
			return self.sentences
		except Exception as e:
			exception = Error( e )
			exception.module = 'processors'
			exception.cause = 'Word'
			exception.method = 'split_sentences( self ) -> List[ str ]'
			Logger( ).write( exception )
			raise exception
	
	def clean_sentences( self ) -> List[ str ] | None:
		"""Clean sentences.
		
		Purpose:
			Executes the ``clean sentences`` operation for the ``WordParser`` workflow, updating
			instance state where required and returning the processed result used by downstream
			Chonky loading, processing, analysis, tokenization, or embedding steps.
		
		Returns:
			List[ str ] | None: Result produced by the processing operation.
		
		Raises:
			Error: Raised when validation or processing fails.
		"""
		try:
			for _sent in self.sentences:
				_sent = re.sub( r'[\r\n\t]+', ' ', _sent )
				_sent = re.sub( r"[^a-zA-Z0-9\s']", '', _sent )
				_sent = re.sub( r'\s{2,}', ' ', _sent ).strip( ).lower( )
				self.cleaned_sentences.append( _sent )
			return self.cleaned_sentences
		except Exception as e:
			exception = Error( e )
			exception.module = 'processors'
			exception.cause = 'Word'
			exception.method = 'clean_sentences( self ) -> List[ str ]'
			Logger( ).write( exception )
			raise exception
	
	def create_vocabulary( self ) -> set | None:
		"""Create vocabulary.
		
		Purpose:
			Executes the ``create vocabulary`` operation for the ``WordParser`` workflow,
			updating instance state where required and returning the processed result used by
			downstream Chonky loading, processing, analysis, tokenization, or embedding steps.
		
		Returns:
			set | None: Result produced by the processing operation.
		
		Raises:
			Error: Raised when validation or processing fails.
		"""
		try:
			all_words = [ ]
			self.stop_words = set( stopwords.words( 'english' ) )
			for _sentence in self.cleaned_sentences:
				_tokens = word_tokenize( _sentence )
				self.tokens = [ token for token in _tokens if
					token.isalpha( ) and token not in self.stop_words ]
				all_words.extend( self.tokens )
			self.vocabulary = set( all_words )
			return self.vocabulary
		except Exception as e:
			exception = Error( e )
			exception.module = 'processors'
			exception.cause = 'Word'
			exception.method = 'create_vocabulary( self ) -> List[ str ]'
			Logger( ).write( exception )
			raise exception
	
	def compute_frequency_distribution( self ) -> Dict[ str, int ] | None:
		"""Compute frequency distribution.
		
		Purpose:
			Executes the ``compute frequency distribution`` operation for the ``WordParser``
			workflow, updating instance state where required and returning the processed result
			used by downstream Chonky loading, processing, analysis, tokenization, or embedding
			steps.
		
		Returns:
			Dict[ str, int ] | None: Result produced by the processing operation.
		
		Raises:
			Error: Raised when validation or processing fails.
		"""
		try:
			words = [ ]
			for sentence in self.cleaned_sentences:
				tokens = word_tokenize( sentence )
				tokens = [ token for token in tokens if token.isalpha( ) ]
				words.extend( tokens )
			self.frequency_distribution = dict( Counter( words ) )
			return self.frequency_distribution
		except Exception as e:
			exception = Error( e )
			exception.module = 'processors'
			exception.cause = 'Word'
			exception.method = 'compute_frequency_distribution( self ) -> Dict[ str, int ]'
			Logger( ).write( exception )
			raise exception
	
	def summarize( self ) -> List[ str ] | None:
		"""Summarize.
		
		Purpose:
			Executes the ``summarize`` operation for the ``WordParser`` workflow, updating
			instance state where required and returning the processed result used by downstream
			Chonky loading, processing, analysis, tokenization, or embedding steps.
		
		Returns:
			List[ str ] | None: Result produced by the processing operation.
		"""
		print( f'Document: {self.file_path}' )
		print( f'Paragraphs: {len( self.paragraphs )}' )
		print( f'Sentences: {len( self.sentences )}' )
		print( f'Vocabulary Size: {len( self.vocabulary )}' )
		print(
			f'Top 10 Frequent Words: {Counter( self.frequency_distribution ).most_common( 10 )}' )

class PdfParser( Processor ):
	"""PdfParser processing component.
	
	Purpose:
		Provides geometry-aware PDF parsing and cleanup utilities that extract page blocks,
		remove repeating headers and footers, repair line spacing, normalize artifacts,
		and rebuild PDF text for downstream processing.
	
	Attributes:
		strip_headers: Runtime state used by ``PdfParser`` during Chonky processing operations.
		minimum_length: Runtime state used by ``PdfParser`` during Chonky processing operations.
		extract_tables_enabled: Runtime state used by ``PdfParser`` during Chonky processing
		operations.
		extracted_lines: Runtime state used by ``PdfParser`` during Chonky processing operations.
		extracted_tables: Runtime state used by ``PdfParser`` during Chonky processing operations.
		extracted_pages: Runtime state used by ``PdfParser`` during Chonky processing operations.
	"""
	strip_headers: Optional[ bool ]
	minimum_length: Optional[ int ]
	extract_tables_enabled: Optional[ bool ]
	extracted_lines: Optional[ List ]
	extracted_tables: Optional[ List ]
	extracted_pages: Optional[ List ]
	
	def __init__( self, headers: bool=False, size: int=10, tables: bool=True ) -> None:
		"""Initialize the PdfParser instance.
		
		Purpose:
			Initializes parser state, reusable helper objects, and runtime caches used by later
			processing methods.
		
		Args:
			headers: Headers value used by the processing operation. Expected type: ``bool``.
			Defaults to ``False``.
			size: Size value used by the processing operation. Expected type: ``int``. Defaults to
			``10``.
			tables: Tables value used by the processing operation. Expected type: ``bool``.
			Defaults to ``True``.
		"""
		super( ).__init__( )
		self.strip_headers = headers
		self.minimum_length = size
		self.extract_tables_enabled = tables
		self.pages = [ ]
		self.lines = [ ]
		self.blocks = [ ]
		self.clean_lines = [ ]
		self.extracted_lines = [ ]
		self.extracted_tables = [ ]
		self.extracted_pages = [ ]
		self.tables = None
		self.file_path = ''
		self.page = ''
	
	def __dir__( self ) -> List[ str ] | None:
		"""Dir.
		
		Purpose:
			Returns the public member names exposed by the parser for introspection, diagnostics,
			and MkDocs API documentation.
		
		Returns:
			List[ str ] | None: Result produced by the processing operation.
		"""
		return [ 'strip_headers', 'minimum_length', 'extract_tables_enabled', 'file_path', 'page',
			'pages', 'lines', 'blocks', 'clean_lines', 'extracted_lines', 'extracted_tables',
			'extracted_pages', 'geometric_extract', 'extract_pages', 'remove_repeats',
			'clean_artifacts', 'repair_spacing', 'rejoin_hyphenation', 'rebuild_pages',
			'extract_lines', 'extract_text', 'extract_tables', 'export_csv', 'export_text',
			'export_excel' ]
	
	def geometric_extract( self, path: str, count: Optional[ int ] = None,
		header_ratio: float=0.08, footer_ratio: float=0.08,
		preserve_page_breaks: bool=False ) -> str | None:
		"""Geometric extract.
		
		Purpose:
			Executes the ``geometric extract`` operation for the ``PdfParser`` workflow, updating
			instance state where required and returning the processed result used by downstream
			Chonky loading, processing, analysis, tokenization, or embedding steps.
		
		Args:
			path: Path value used by the processing operation. Expected type: ``str``.
			count: Count value used by the processing operation. Expected type: ``Optional[ int
			]``. Defaults to ``None``.
			header_ratio: Header ratio value used by the processing operation. Expected type:
			``float``. Defaults to ``0.08``.
			footer_ratio: Footer ratio value used by the processing operation. Expected type:
			``float``. Defaults to ``0.08``.
			preserve_page_breaks: Preserve page breaks value used by the processing operation.
			Expected type: ``bool``. Defaults to ``False``.
		
		Returns:
			str | None: Result produced by the processing operation.
		
		Raises:
			Error: Raised when validation or processing fails.
		"""
		try:
			throw_if( 'path', path )
			pages = self.extract_pages( path=path, count=count, header_ratio=header_ratio,
				footer_ratio=footer_ratio ) or [ ]
			
			return self.rebuild_pages( pages=pages, preserve_page_breaks=preserve_page_breaks )
		except Exception as e:
			exception = Error( e )
			exception.module = 'processors'
			exception.cause = 'PdfParser'
			exception.method = 'geometric_extract( self, path: str ) -> str'
			Logger( ).write( exception )
			raise exception
	
	def extract_pages( self, path: str, count: Optional[ int ] = None, header_ratio: float=0.08,
		footer_ratio: float=0.08 ) -> List[ dict ] | None:
		"""Extract pages.
		
		Purpose:
			Executes the ``extract pages`` operation for the ``PdfParser`` workflow, updating
			instance state where required and returning the processed result used by downstream
			Chonky loading, processing, analysis, tokenization, or embedding steps.
		
		Args:
			path: Path value used by the processing operation. Expected type: ``str``.
			count: Count value used by the processing operation. Expected type: ``Optional[ int
			]``. Defaults to ``None``.
			header_ratio: Header ratio value used by the processing operation. Expected type:
			``float``. Defaults to ``0.08``.
			footer_ratio: Footer ratio value used by the processing operation. Expected type:
			``float``. Defaults to ``0.08``.
		
		Returns:
			List[ dict ] | None: Result produced by the processing operation.
		
		Raises:
			Error: Raised when validation or processing fails.
		"""
		try:
			throw_if( 'path', path )
			
			if header_ratio < 0.0 or header_ratio > 0.30:
				raise ValueError( 'Argument "header_ratio" must be between 0.00 and 0.30.' )
			
			if footer_ratio < 0.0 or footer_ratio > 0.30:
				raise ValueError( 'Argument "footer_ratio" must be between 0.00 and 0.30.' )
			
			self.file_path = path
			self.pages = [ ]
			with fitz.open( self.file_path ) as doc:
				for page_index, page in enumerate( doc ):
					if count is not None and page_index >= count:
						break
					
					page_height = float( page.rect.height )
					header_limit = page_height * header_ratio
					footer_limit = page_height * (1.0 - footer_ratio)
					raw_blocks = page.get_text( 'blocks' ) or [ ]
					page_blocks = [ ]
					for block_index, block in enumerate( raw_blocks ):
						if len( block ) < 5:
							continue
						
						x0 = float( block[ 0 ] )
						y0 = float( block[ 1 ] )
						x1 = float( block[ 2 ] )
						y1 = float( block[ 3 ] )
						text = block[ 4 ]
						block_type = block[ 6 ] if len( block ) > 6 else 0
						if block_type != 0:
							continue
						
						if not isinstance( text, str ) or not text.strip( ):
							continue
						
						midpoint = (y0 + y1) / 2.0
						zone = 'body'
						if midpoint <= header_limit:
							zone = 'header'
						elif midpoint >= footer_limit:
							zone = 'footer'
						
						page_blocks.append( { 'page': page_index + 1, 'index': block_index, 'x0': x0, 'y0': y0,
								'x1': x1, 'y1': y1, 'midpoint': midpoint, 'zone': zone,
								'text': text, 'drop': False, } )
					
					page_blocks.sort( key=lambda item: (item[ 'y0' ], item[ 'x0' ]) )
					self.pages.append( { 'page': page_index + 1, 'width': float( page.rect.width ),
						'height': page_height, 'blocks': page_blocks, } )
			
			return self.pages
		except Exception as e:
			exception = Error( e )
			exception.module = 'processors'
			exception.cause = 'PdfParser'
			exception.method = 'extract_pages( self, path: str ) -> List[ dict ]'
			Logger( ).write( exception )
			raise exception
	
	def remove_repeats( self, pages: List[ Dict ], minimum_repeats: int=3 ) -> List[ Dict ] | None:
		"""Remove repeats.
		
		Purpose:
			Executes the ``remove repeats`` operation for the ``PdfParser`` workflow, updating
			instance state where required and returning the processed result used by downstream
			Chonky loading, processing, analysis, tokenization, or embedding steps.
		
		Args:
			pages: Pages value used by the processing operation. Expected type: ``List[ dict ]``.
			minimum_repeats: Minimum repeats value used by the processing operation. Expected
			type: ``int``. Defaults to ``3``.
		
		Returns:
			List[
				                                                                             dict
				                                                                             ] |
				                                                                             None:
				                                                                             Result produced by the processing operation.
		
		Raises:
			Error: Raised when validation or processing fails.
		"""
		try:
			throw_if( 'pages', pages )
			
			def normalize_candidate( value: str ) -> str:
				candidate = value.strip( ).lower( )
				candidate = re.sub( r'\s+', ' ', candidate )
				candidate = re.sub( r'\b\d{1,5}\b', '#', candidate )
				candidate = re.sub( r'\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b', '#date#', candidate )
				candidate = re.sub( r'\s+', ' ', candidate ).strip( )
				return candidate
			
			def is_page_label( value: str ) -> bool:
				candidate = value.strip( )
				
				if re.fullmatch( r'\d{1,5}', candidate ):
					return True
				
				if re.fullmatch( r'[ivxlcdm]{1,10}', candidate, flags=re.IGNORECASE ):
					return True
				
				if re.fullmatch( r'(?:page|p\.)\s+\d+(?:\s+of\s+\d+)?', candidate,
						flags=re.IGNORECASE ):
					return True
				
				return False
			
			counts = { }
			threshold = max( 2, int( minimum_repeats ) )
			
			for page in pages:
				for block in page.get( 'blocks', [ ] ):
					if block.get( 'zone' ) not in { 'header', 'footer' }:
						continue
					
					text = block.get( 'text' )
					
					if not isinstance( text, str ) or not text.strip( ):
						continue
					
					key = normalize_candidate( text )
					
					if key:
						counts[ key ] = counts.get( key, 0 ) + 1
			
			repeated = { key for key, value in counts.items( ) if value >= threshold }
			cleaned_pages = [ ]
			
			for page in pages:
				page_copy = dict( page )
				new_blocks = [ ]
				
				for block in page.get( 'blocks', [ ] ):
					block_copy = dict( block )
					text = block_copy.get( 'text' )
					
					if isinstance( text, str ) and block_copy.get( 'zone' ) in { 'header',
						'footer' }:
						key = normalize_candidate( text )
						
						if key in repeated or is_page_label( text ):
							block_copy[ 'drop' ] = True
					
					new_blocks.append( block_copy )
				
				page_copy[ 'blocks' ] = new_blocks
				cleaned_pages.append( page_copy )
			
			return cleaned_pages
		except Exception as e:
			exception = Error( e )
			exception.module = 'processors'
			exception.cause = 'PdfParser'
			exception.method = 'remove_repeats( self, pages: List[ dict ] )'
			Logger( ).write( exception )
			raise exception
	
	def clean_artifacts( self, text: str ) -> str:
		"""Clean artifacts.
		
		Purpose:
			Executes the ``clean artifacts`` operation for the ``PdfParser`` workflow, updating
			instance state where required and returning the processed result used by downstream
			Chonky loading, processing, analysis, tokenization, or embedding steps.
		
		Args:
			text: Text value used by the processing operation. Expected type: ``str``.
		
		Returns:
			str: Result produced by the processing operation.
		
		Raises:
			Error: Raised when validation or processing fails.
		"""
		try:
			throw_if( 'text', text )
			cleaned = text
			
			cleaned = re.sub( r'<parsed\s+text\s+for\s+page:\s*\d+\s*/\s*\d+>', ' ', cleaned,
				flags=re.IGNORECASE )
			
			cleaned = re.sub( r'<image\s+for\s+page:\s*\d+\s*/\s*\d+>', ' ', cleaned,
				flags=re.IGNORECASE )
			
			cleaned = re.sub( r'</?[a-z][a-z0-9_-]{1,20}>', ' ', cleaned, flags=re.IGNORECASE )
			
			cleaned = re.sub( r'\b[a-z]:\\[^\s<>]*(?:\.[a-z0-9]{2,5})\b', ' ', cleaned,
				flags=re.IGNORECASE )
			
			cleaned = re.sub( r'\b(?:/[^/\s<>]+)+/(?:[^/\s<>]+\.[a-z0-9]{2,5})\b', ' ', cleaned,
				flags=re.IGNORECASE )
			
			cleaned = re.sub( r'(?im)^\s*(?:endobj|obj|xref|trailer|startxref|%%eof)\s*$', ' ',
				cleaned )
			
			cleaned = re.sub( r'(?im)^\s*[a-z][a-z0-9_.-]{1,40}\s+on\s+[a-z0-9_.-]{6,}'
			                  r'(?:\s+with\s+[a-z0-9_.-]+)?\s*$', ' ', cleaned )
			
			cleaned = re.sub( r'(?<=\S)\s*\.{4,}\s*(?=\S)', ' ', cleaned )
			cleaned = re.sub( r'(?<=\S)\s*(?:\.\s*){4,}(?=\S)', ' ', cleaned )
			cleaned = re.sub( r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]', '', cleaned )
			cleaned = re.sub( r'[ \t]{2,}', ' ', cleaned )
			cleaned = re.sub( r' +\n', '\n', cleaned )
			cleaned = re.sub( r'\n +', '\n', cleaned )
			cleaned = re.sub( r'\n{3,}', '\n\n', cleaned )
			
			return cleaned.strip( )
		except Exception as e:
			exception = Error( e )
			exception.module = 'processors'
			exception.cause = 'PdfParser'
			exception.method = 'clean_artifacts( self, text: str ) -> str'
			Logger( ).write( exception )
			raise exception
	
	def repair_spacing( self, text: str ) -> str:
		"""Repair spacing.
		
		Purpose:
			Repairs malformed punctuation spacing and extracted letter-spaced text while preserving
			valid URL schemes, legal citations, capitalization, line boundaries, and document
			order.
		
		Args:
			text: Source text containing spacing artifacts produced during PDF extraction.
		
		Returns:
			str: Text with supported spacing artifacts repaired.
		
		Raises:
			Error: Raised when validation or spacing repair fails.
		"""
		try:
			throw_if( 'text', text )
			cleaned = text
			def repair_bracketed( match: re.Match ) -> str:
				value = match.group( 1 )
				
				if re.fullmatch( r'[A-Z\s\-\n\r]{3,}', value ):
					value = re.sub( r'[\s\-]+', '', value )
					return f'[{value}]'
				
				return match.group( 0 )
			
			def repair_letter_spaced_line( match: re.Match ) -> str:
				line = match.group( 0 )
				tokens = line.split( )
				single_letters = [ token for token in tokens if re.fullmatch( r'[A-Z]', token ) ]
				if (len( single_letters ) >= 3 and len( single_letters ) >= len( tokens ) / 2):
					return re.sub( r'\b([A-Z])(?:\s+)(?=[A-Z]\b)', r'\1', line )
				
				return line
			
			cleaned = re.sub( r'\[([A-Z\s\-\n\r]{3,})]', repair_bracketed, cleaned )
			
			cleaned = re.sub( r'(?m)^[A-Z](?:\s+[A-Z]){2,}(?:\s+[A-Z]{2,})*\s*$',
				repair_letter_spaced_line, cleaned )
			
			cleaned = re.sub( r'[ \t]+([,.;:!?])', r'\1', cleaned )
			cleaned = re.sub( r'([!?;])(?=\S)', r'\1 ', cleaned )
			cleaned = re.sub( r':(?!//)(?=\S)', ': ', cleaned )
			cleaned = re.sub( r'(?<=[a-z0-9])\.(?=[A-Z][a-z])', '. ', cleaned )
			cleaned = re.sub( r'[ \t]{2,}', ' ', cleaned )
			cleaned = re.sub( r' +\n', '\n', cleaned )
			cleaned = re.sub( r'\n +', '\n', cleaned )
			cleaned = re.sub( r'\n{3,}', '\n\n', cleaned )
			
			return cleaned.strip( )
		except Exception as e:
			exception = Error( e )
			exception.module = 'processors'
			exception.cause = 'PdfParser'
			exception.method = 'repair_spacing( self, text: str ) -> str'
			Logger( ).write( exception )
			raise exception
	
	def rejoin_hyphenation( self, text: str, repair_embedded: bool=True ) -> str:
		"""Rejoin hyphenation.
		
		Purpose:
			Executes the ``rejoin hyphenation`` operation for the ``PdfParser`` workflow,
			updating instance state where required and returning the processed result used by
			downstream Chonky loading, processing, analysis, tokenization, or embedding steps.
		
		Args:
			text: Text value used by the processing operation. Expected type: ``str``.
			repair_embedded: Repair embedded value used by the processing operation. Expected
			type: ``bool``. Defaults to ``True``.
		
		Returns:
			str: Result produced by the processing operation.
		
		Raises:
			Error: Raised when validation or processing fails.
		"""
		try:
			throw_if( 'text', text )
			cleaned = text
			try:
				from nltk.corpus import words as nltk_words
				
				vocabulary = set( nltk_words.words( 'en' ) )
			except Exception:
				vocabulary = set( )
			
			try:
				from nltk.corpus import wordnet as nltk_wordnet
			except Exception:
				nltk_wordnet = None
			
			def is_known_word( value: str ) -> bool:
				token = value.lower( )
				
				if token in vocabulary:
					return True
				
				if nltk_wordnet is not None:
					try:
						if nltk_wordnet.synsets( token ):
							return True
						
						if nltk_wordnet.morphy( token ) is not None:
							return True
					except Exception:
						return False
				
				return False
			
			def repair_line_break( match: re.Match ) -> str:
				left = match.group( 1 )
				right = match.group( 2 )
				
				if '-' in right:
					return f'{left}-{right}'
				
				return f'{left}{right}'
			
			def repair_embedded_split( match: re.Match ) -> str:
				left = match.group( 1 )
				right = match.group( 2 )
				combined = f'{left}{right}'
				
				combined_known = is_known_word( combined )
				
				if combined_known:
					return combined
				
				if re.fullmatch( r'(?:able|ible|ally|ance|ancy|ence|ency|ation|ations|'
				                 r'cation|cations|cies|tion|tions|sion|sions|ment|ments|'
				                 r'ness|less|ship|ships|ing|ings|ed|er|ers|or|ors|ies|'
				                 r'ive|ives|ity|ities|al|als|ary|ory|ories|ous|ious|eous)', right,
						flags=re.IGNORECASE ):
					return combined
				
				return match.group( 0 )
			
			cleaned = re.sub( r'(?<=[A-Za-z])[\u00AD\uFFFC\uFFFD\uFFFE]\s*(?=[A-Za-z])', '',
				cleaned )
			
			cleaned = re.sub( r'\b([A-Za-z]{2,})-\s*\n\s*([A-Za-z][A-Za-z-]*)\b', 
				repair_line_break,
				cleaned )
			
			cleaned = re.sub( r'\b([A-Za-z]{2,})-\s+([A-Za-z][A-Za-z-]*)\b', repair_line_break,
				cleaned )
			
			if repair_embedded:
				cleaned = re.sub( r'\b([A-Za-z]{2,})-([a-z]{2,})\b', repair_embedded_split,
					cleaned )
			
			cleaned = re.sub( r'[ \t]{2,}', ' ', cleaned )
			cleaned = re.sub( r' +\n', '\n', cleaned )
			cleaned = re.sub( r'\n +', '\n', cleaned )
			cleaned = re.sub( r'\n{3,}', '\n\n', cleaned )
			
			return cleaned.strip( )
		except Exception as e:
			exception = Error( e )
			exception.module = 'processors'
			exception.cause = 'PdfParser'
			exception.method = 'rejoin_hyphenation( self, text: str ) -> str'
			Logger( ).write( exception )
			raise exception
	
	def rebuild_pages( self, pages: List[ dict ], preserve_page_breaks: bool=False ) -> str:
		"""Rebuild pages.
		
		Purpose:
			Executes the ``rebuild pages`` operation for the ``PdfParser`` workflow, updating
			instance state where required and returning the processed result used by downstream
			Chonky loading, processing, analysis, tokenization, or embedding steps.
		
		Args:
			pages: Pages value used by the processing operation. Expected type: ``List[ dict ]``.
			preserve_page_breaks: Preserve page breaks value used by the processing operation.
			Expected type: ``bool``. Defaults to ``False``.
		
		Returns:
			str: Result produced by the processing operation.
		
		Raises:
			Error: Raised when validation or processing fails.
		"""
		try:
			throw_if( 'pages', pages )
			page_texts = [ ]
			
			for page in pages:
				page_number = page.get( 'page' )
				blocks = page.get( 'blocks', [ ] )
				parts = [ ]
				
				for block in blocks:
					if block.get( 'drop' ) is True:
						continue
					
					text = block.get( 'text' )
					
					if not isinstance( text, str ) or not text.strip( ):
						continue
					
					block_lines = [ line.strip( ) for line in text.splitlines( ) if
						isinstance( line, str ) and line.strip( ) ]
					
					if block_lines:
						parts.append( '\n'.join( block_lines ) )
				
				page_text = '\n\n'.join( parts )
				page_text = re.sub( r'[ \t]{2,}', ' ', page_text )
				page_text = re.sub( r' +\n', '\n', page_text )
				page_text = re.sub( r'\n +', '\n', page_text )
				page_text = re.sub( r'\n{3,}', '\n\n', page_text ).strip( )
				
				if not page_text:
					continue
				
				if preserve_page_breaks:
					page_texts.append( f'<<<PAGE_BREAK:{page_number}>>>\n{page_text}' )
				else:
					page_texts.append( page_text )
			
			return '\n\n'.join( page_texts ).strip( )
		except Exception as e:
			exception = Error( e )
			exception.module = 'processors'
			exception.cause = 'PdfParser'
			exception.method = 'rebuild_pages( self, pages: List[ dict ] ) -> str'
			Logger( ).write( exception )
			raise exception
	
	def extract_lines( self, path: str, count: Optional[ int ] = None ) -> List[ str ] | None:
		"""Extract lines.
		
		Purpose:
			Executes the ``extract lines`` operation for the ``PdfParser`` workflow, updating
			instance state where required and returning the processed result used by downstream
			Chonky loading, processing, analysis, tokenization, or embedding steps.
		
		Args:
			path: Path value used by the processing operation. Expected type: ``str``.
			count: Count value used by the processing operation. Expected type: ``Optional[ int
			]``. Defaults to ``None``.
		
		Returns:
			List[ str ] | None: Result produced by the processing operation.
		
		Raises:
			Error: Raised when validation or processing fails.
		"""
		try:
			throw_if( 'path', path )
			text = self.geometric_extract( path=path, count=count ) or ''
			self.extracted_lines = [ line.strip( ) for line in text.splitlines( ) if
				isinstance( line, str ) and line.strip( ) ]
			return self.extracted_lines
		except Exception as e:
			exception = Error( e )
			exception.module = 'processors'
			exception.cause = 'PdfParser'
			exception.method = 'extract_lines( self, path: str, count: Optional[ int ]=None )'
			Logger( ).write( exception )
			raise exception
	
	def extract_text( self, path: str, count: Optional[ int ] = None ) -> str | None:
		"""Extract text.
		
		Purpose:
			Executes the ``extract text`` operation for the ``PdfParser`` workflow, updating
			instance state where required and returning the processed result used by downstream
			Chonky loading, processing, analysis, tokenization, or embedding steps.
		
		Args:
			path: Path value used by the processing operation. Expected type: ``str``.
			count: Count value used by the processing operation. Expected type: ``Optional[ int
			]``. Defaults to ``None``.
		
		Returns:
			str | None: Result produced by the processing operation.
		
		Raises:
			Error: Raised when validation or processing fails.
		"""
		try:
			throw_if( 'path', path )
			return self.geometric_extract( path=path, count=count )
		except Exception as e:
			exception = Error( e )
			exception.module = 'processors'
			exception.cause = 'PdfParser'
			exception.method = 'extract_text( self, path: str, count: Optional[ int ]=None )'
			Logger( ).write( exception )
			raise exception
	
	def extract_tables( self, path: str, count: Optional[ int ] = None ) -> ( List[ pd.DataFrame ] | None):
		"""Extract tables.
		
		Purpose:
			Executes the ``extract tables`` operation for the ``PdfParser`` workflow, updating
			instance state where required and returning the processed result used by downstream
			Chonky loading, processing, analysis, tokenization, or embedding steps.
		
		Args:
			path: Path value used by the processing operation. Expected type: ``str``.
			count: Count value used by the processing operation. Expected type: ``Optional[ int
			]``. Defaults to ``None``.
		
		Returns:
			List[
				                                                                        pd.DataFrame ] | None: Result produced by the processing operation.
		
		Raises:
			Error: Raised when validation or processing fails.
		"""
		try:
			throw_if( 'path', path )
			self.file_path = path
			self.tables = [ ]
			with fitz.open( self.file_path ) as doc:
				for page_index, page in enumerate( doc ):
					if count is not None and page_index >= count:
						break
					
					tables = page.find_tables( )
					for table in getattr( tables, 'tables', [ ] ):
						self.tables.append( pd.DataFrame( table.extract( ) ) )
			
			return self.tables
		except Exception as e:
			exception = Error( e )
			exception.module = 'processors'
			exception.cause = 'PdfParser'
			exception.method = 'extract_tables( self, path: str, count: Optional[ int ]=None )'
			Logger( ).write( exception )
			raise exception
	
	def export_csv( self, tables: List[ pd.DataFrame ], filename: str ) -> None:
		"""Export csv.
		
		Purpose:
			Executes the ``export csv`` operation for the ``PdfParser`` workflow, updating
			instance state where required and returning the processed result used by downstream
			Chonky loading, processing, analysis, tokenization, or embedding steps.
		
		Args:
			tables: Tables value used by the processing operation. Expected type: ``List[
			pd.DataFrame ]``.
			filename: Filename value used by the processing operation. Expected type: ``str``.
		
		Returns:
			None: Result produced by the processing operation.
		
		Raises:
			Error: Raised when validation or processing fails.
		"""
		try:
			throw_if( 'tables', tables )
			throw_if( 'filename', filename )
			self.tables = tables
			
			for index, df_table in enumerate( self.tables ):
				df_table.to_csv( f'{filename}_{index + 1}.csv', index=False )
		except Exception as e:
			exception = Error( e )
			exception.module = 'processors'
			exception.cause = 'PdfParser'
			exception.method = 'export_csv( self, tables: List[ pd.DataFrame ], filename: str )'
			Logger( ).write( exception )
			raise exception
	
	def export_text( self, lines: List[ str ], path: str ) -> None:
		"""Export text.
		
		Purpose:
			Executes the ``export text`` operation for the ``PdfParser`` workflow, updating
			instance state where required and returning the processed result used by downstream
			Chonky loading, processing, analysis, tokenization, or embedding steps.
		
		Args:
			lines: Lines value used by the processing operation. Expected type: ``List[ str ]``.
			path: Path value used by the processing operation. Expected type: ``str``.
		
		Returns:
			None: Result produced by the processing operation.
		
		Raises:
			Error: Raised when validation or processing fails.
		"""
		try:
			throw_if( 'lines', lines )
			throw_if( 'path', path )
			self.lines = lines
			self.file_path = path
			
			with open( self.file_path, 'w', encoding='utf-8', errors='ignore' ) as file:
				for line in self.lines:
					file.write( line + '\n' )
		except Exception as e:
			exception = Error( e )
			exception.module = 'processors'
			exception.cause = 'PdfParser'
			exception.method = 'export_text( self, lines: List[ str ], path: str )'
			Logger( ).write( exception )
			raise exception
	
	def export_excel( self, tables: List[ pd.DataFrame ], path: str ) -> None:
		"""Export excel.
		
		Purpose:
			Executes the ``export excel`` operation for the ``PdfParser`` workflow, updating
			instance state where required and returning the processed result used by downstream
			Chonky loading, processing, analysis, tokenization, or embedding steps.
		
		Args:
			tables: Tables value used by the processing operation. Expected type: ``List[
			pd.DataFrame ]``.
			path: Path value used by the processing operation. Expected type: ``str``.
		
		Returns:
			None: Result produced by the processing operation.
		
		Raises:
			Error: Raised when validation or processing fails.
		"""
		try:
			throw_if( 'tables', tables )
			throw_if( 'path', path )
			self.tables = tables
			self.file_path = path
			
			with pd.ExcelWriter( self.file_path, engine='xlsxwriter' ) as writer:
				for index, df_table in enumerate( self.tables ):
					sheet = f'Table_{index + 1}'
					df_table.to_excel( writer, sheet_name=sheet, index=False )
		except Exception as e:
			exception = Error( e )
			exception.module = 'processors'
			exception.cause = 'PdfParser'
			exception.method = 'export_excel( self, tables: List[ pd.DataFrame ], path: str )'
			Logger( ).write( exception )
			raise exception
